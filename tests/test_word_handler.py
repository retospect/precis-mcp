"""Integration tests for WordHandler — parse, read, put."""

from __future__ import annotations

import pytest
from docx import Document

from precis.handlers.word import WordHandler
from precis.protocol import PrecisError


@pytest.fixture
def handler():
    return WordHandler()


@pytest.fixture
def sample_docx(tmp_path):
    """Create a sample DOCX with headings and paragraphs."""
    path = tmp_path / "test.docx"
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("This is the first paragraph of the introduction.")
    doc.add_paragraph("This is the second paragraph.")
    doc.add_heading("Methods", level=1)
    doc.add_paragraph("We used a novel approach to solve the problem.")
    doc.add_heading("Results", level=1)
    doc.add_heading("Subsection A", level=2)
    doc.add_paragraph("The results show significant improvement.")
    doc.save(str(path))
    return path


@pytest.fixture
def empty_docx(tmp_path):
    """Create an empty DOCX."""
    path = tmp_path / "empty.docx"
    doc = Document()
    doc.save(str(path))
    return path


# ── Parse tests ─────────────────────────────────────────────────────


class TestParse:
    def test_parse_returns_nodes(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        assert len(nodes) > 0

    def test_parse_headings(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        headings = [n for n in nodes if n.node_type == "h"]
        assert len(headings) == 4
        assert headings[0].text == "Introduction"
        assert headings[1].text == "Methods"
        assert headings[2].text == "Results"
        assert headings[3].text == "Subsection A"

    def test_parse_paragraphs(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        paras = [n for n in nodes if n.node_type == "p"]
        assert len(paras) == 4

    def test_parse_heading_levels(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        headings = [n for n in nodes if n.node_type == "h"]
        assert headings[0].heading_level() == 1
        assert headings[3].heading_level() == 2

    def test_parse_paths(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        headings = [n for n in nodes if n.node_type == "h"]
        assert str(headings[0].path) == "S1"
        assert str(headings[1].path) == "S2"
        assert str(headings[2].path) == "S3"
        assert str(headings[3].path) == "S3.1"

    def test_parse_slugs_unique(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        slugs = [n.slug for n in nodes]
        assert len(slugs) == len(set(slugs))

    def test_parse_empty_doc(self, handler, empty_docx):
        nodes = handler.parse(empty_docx)
        assert len(nodes) == 0

    def test_source_files(self, handler, sample_docx):
        files = handler.source_files(sample_docx)
        assert len(files) == 1
        assert files[0] == sample_docx


# ── Read tests ──────────────────────────────────────────────────────


class TestRead:
    def test_read_toc(self, handler, sample_docx):
        result = handler.read(str(sample_docx), None, None, None, "", False, 0, 1)
        assert "test.docx" in result
        assert "Introduction" in result
        assert "Methods" in result

    def test_read_toc_depth(self, handler, sample_docx):
        result = handler.read(str(sample_docx), None, None, None, "", False, 1, 1)
        # depth=1 should show only level-1 headings
        assert "Introduction" in result
        assert "Subsection A" not in result

    def test_read_selector_heading(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        intro = [n for n in nodes if n.text == "Introduction"][0]
        result = handler.read(str(sample_docx), intro.slug, None, None, "", False, 0, 1)
        assert "Introduction" in result
        # Children shown as precis lines (multi-node section)
        para_slugs = [n.slug for n in nodes if n.node_type == "p" and n.path.h1 == 1]
        for slug in para_slugs:
            assert slug in result

    def test_read_selector_paragraph(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        para = [n for n in nodes if n.node_type == "p"][0]
        result = handler.read(str(sample_docx), para.slug, None, None, "", False, 0, 1)
        assert para.text in result

    def test_read_query(self, handler, sample_docx):
        result = handler.read(
            str(sample_docx), None, None, None, "novel approach", False, 0, 1
        )
        assert (
            "novel approach" in result.lower() or "1 hits" in result or "hit" in result
        )

    def test_read_query_no_match(self, handler, sample_docx):
        result = handler.read(
            str(sample_docx), None, None, None, "xyznonexistent", False, 0, 1
        )
        assert "No matches" in result or "0 hits" in result

    def test_read_meta(self, handler, sample_docx):
        result = handler.read(str(sample_docx), None, "meta", None, "", False, 0, 1)
        assert "nodes:" in result
        assert "headings:" in result

    def test_read_empty_doc(self, handler, empty_docx):
        result = handler.read(str(empty_docx), None, None, None, "", False, 0, 1)
        assert "empty" in result.lower() or "0 nodes" in result


# ── Write tests ─────────────────────────────────────────────────────


class TestPut:
    def test_put_append(self, handler, sample_docx):
        result = handler.put(str(sample_docx), None, "A new paragraph.", "append")
        assert "+" in result
        # Verify it was written
        nodes = handler.parse(sample_docx)
        texts = [n.text for n in nodes]
        assert "A new paragraph." in texts

    def test_put_append_heading(self, handler, sample_docx):
        result = handler.put(str(sample_docx), None, "# | Discussion", "append")
        assert "+" in result
        nodes = handler.parse(sample_docx)
        headings = [n for n in nodes if n.node_type == "h"]
        assert any(h.text == "Discussion" for h in headings)

    def test_put_replace(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        para = [n for n in nodes if n.node_type == "p"][0]
        result = handler.put(
            str(sample_docx),
            para.slug,
            "Replaced text here.",
            "replace",
            tracked=False,
        )
        assert "replace" in result.lower()
        new_nodes = handler.parse(sample_docx)
        new_texts = [n.text for n in new_nodes if n.node_type == "p"]
        assert "Replaced text here." in new_texts

    def test_put_delete(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        para = [n for n in nodes if n.node_type == "p"][0]
        count_before = len(nodes)
        handler.put(str(sample_docx), para.slug, "", "delete")
        new_nodes = handler.parse(sample_docx)
        assert len(new_nodes) == count_before - 1

    def test_put_insert_after(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        para = [n for n in nodes if n.node_type == "p"][0]
        result = handler.put(str(sample_docx), para.slug, "Inserted after.", "after")
        assert "+" in result
        new_nodes = handler.parse(sample_docx)
        texts = [n.text for n in new_nodes]
        assert "Inserted after." in texts

    def test_put_insert_before(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        para = [n for n in nodes if n.node_type == "p"][0]
        result = handler.put(str(sample_docx), para.slug, "Inserted before.", "before")
        assert "+" in result
        new_nodes = handler.parse(sample_docx)
        texts = [n.text for n in new_nodes]
        assert "Inserted before." in texts

    def test_put_invalid_mode(self, handler, sample_docx):
        with pytest.raises(PrecisError, match="invalid mode"):
            handler.put(str(sample_docx), None, "text", "badmode")

    def test_put_replace_no_text(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        para = [n for n in nodes if n.node_type == "p"][0]
        with pytest.raises(PrecisError, match="text required"):
            handler.put(str(sample_docx), para.slug, "", "replace")

    def test_put_append_no_text(self, handler, sample_docx):
        with pytest.raises(PrecisError, match="text required"):
            handler.put(str(sample_docx), None, "", "append")


# ── Track changes tests ─────────────────────────────────────────────


class TestTrackedChanges:
    def test_write_tracked(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        para = [n for n in nodes if n.node_type == "p"][0]
        result = handler.put(
            str(sample_docx),
            para.slug,
            "Tracked replacement.",
            "replace",
            tracked=True,
        )
        assert "tracked" in result.lower() or "replace" in result.lower()
        # Doc should still be valid
        doc = Document(str(sample_docx))
        assert doc is not None


# ── Comment tests ───────────────────────────────────────────────────


class TestComments:
    def test_write_comment(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        para = [n for n in nodes if n.node_type == "p"][0]
        comment_id = handler.write_comment(sample_docx, para, "Test comment", "tester")
        assert comment_id >= 1

        # Verify comment appears in parse
        new_nodes = handler.parse(sample_docx)
        commented = [n for n in new_nodes if n.comments]
        assert len(commented) >= 1
        assert any(c["text"] == "Test comment" for n in commented for c in n.comments)

    def test_put_note_mode(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        para = [n for n in nodes if n.node_type == "p"][0]
        result = handler.put(str(sample_docx), para.slug, "Note via put", "note")
        assert "comment" in result.lower() or "💬" in result


# ── Table tests ─────────────────────────────────────────────────────


class TestTables:
    def test_parse_table(self, handler, tmp_path):
        path = tmp_path / "table.docx"
        doc = Document()
        doc.add_heading("Data", level=1)
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "Name"
        table.rows[0].cells[1].text = "Value"
        table.rows[1].cells[0].text = "Alpha"
        table.rows[1].cells[1].text = "1"
        table.rows[2].cells[0].text = "Beta"
        table.rows[2].cells[1].text = "2"
        doc.save(str(path))

        nodes = handler.parse(path)
        tables = [n for n in nodes if n.node_type == "t"]
        assert len(tables) == 1
        assert "Name" in tables[0].text
        assert "Value" in tables[0].text


# ── Move tests ──────────────────────────────────────────────────────


class TestMove:
    def test_move_node(self, handler, sample_docx):
        nodes = handler.parse(sample_docx)
        paras = [n for n in nodes if n.node_type == "p"]
        assert len(paras) >= 2

        # Move second paragraph after the first
        result = handler.put(str(sample_docx), paras[1].slug, paras[0].slug, "move")
        assert "moved" in result.lower()


# ── Auto-create tests ──────────────────────────────────────────────


class TestLists:
    def test_parse_bullet_list(self, handler, tmp_path):
        """Bullet list paragraphs should be parsed with '- ' prefix."""
        path = tmp_path / "lists.docx"
        doc = Document()
        doc.add_heading("Shopping", level=1)
        doc.add_paragraph("Apples", style="List Bullet")
        doc.add_paragraph("Bananas", style="List Bullet")
        doc.save(str(path))

        nodes = handler.parse(path)
        paras = [n for n in nodes if n.node_type == "p"]
        assert len(paras) == 2
        assert paras[0].text.startswith("- ")
        assert "Apples" in paras[0].text
        assert paras[1].text.startswith("- ")
        assert "Bananas" in paras[1].text

    def test_parse_numbered_list(self, handler, tmp_path):
        """Numbered list paragraphs should be parsed with '1. ' prefix."""
        path = tmp_path / "numlist.docx"
        doc = Document()
        doc.add_heading("Steps", level=1)
        doc.add_paragraph("First step", style="List Number")
        doc.add_paragraph("Second step", style="List Number")
        doc.save(str(path))

        nodes = handler.parse(path)
        paras = [n for n in nodes if n.node_type == "p"]
        assert len(paras) == 2
        assert "1. " in paras[0].text
        assert "First step" in paras[0].text

    def test_append_bullet_list(self, handler, sample_docx):
        """Appending '- item' text should create list-styled paragraphs."""
        result = handler.put(
            str(sample_docx),
            None,
            "- Alpha\n- Beta\n- Gamma",
            "append",
        )
        assert "+" in result
        nodes = handler.parse(sample_docx)
        list_nodes = [n for n in nodes if n.text.startswith("- ")]
        assert len(list_nodes) >= 3

    def test_append_numbered_list(self, handler, sample_docx):
        """Appending '1. item' text should create numbered list paragraphs."""
        result = handler.put(
            str(sample_docx),
            None,
            "1. First\n2. Second",
            "append",
        )
        assert "+" in result
        nodes = handler.parse(sample_docx)
        num_nodes = [n for n in nodes if "1. " in n.text]
        assert len(num_nodes) >= 2

    def test_list_roundtrip(self, handler, tmp_path):
        """Write bullet list → parse → verify prefix preserved."""
        path = tmp_path / "rt.docx"
        doc = Document()
        doc.add_heading("List", level=1)
        doc.save(str(path))

        handler.put(str(path), None, "- Item A\n- Item B", "append")
        nodes = handler.parse(path)
        items = [n for n in nodes if n.text.startswith("- ")]
        assert len(items) == 2
        assert "Item A" in items[0].text
        assert "Item B" in items[1].text


class TestAutoCreate:
    def test_auto_create_docx(self, handler, tmp_path):
        path = tmp_path / "new.docx"
        assert not path.exists()
        result = handler.read(str(path), None, None, None, "", False, 0, 1)
        assert path.exists()
        assert "0 nodes" in result or "empty" in result.lower()
