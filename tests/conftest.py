"""Shared fixtures for precis tests."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document


@pytest.fixture
def tmp_docx(tmp_path: Path) -> Path:
    """Create a simple test DOCX with headings, paragraphs, and a table."""
    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("We present a novel approach to wibble detection.")
    doc.add_paragraph("Related work is sparse with two prior attempts.")
    doc.add_heading("Methods", level=2)
    doc.add_paragraph("We employ a 12-layer transformer architecture.")

    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Param"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Layers"
    table.rows[1].cells[1].text = "12"
    table.rows[2].cells[0].text = "Hidden"
    table.rows[2].cells[1].text = "768"

    p = tmp_path / "test.docx"
    doc.save(str(p))
    return p


@pytest.fixture
def empty_docx(tmp_path: Path) -> Path:
    """Create an empty DOCX."""
    doc = Document()
    p = tmp_path / "empty.docx"
    doc.save(str(p))
    return p


@pytest.fixture
def formatted_docx(tmp_path: Path) -> Path:
    """Create a DOCX with formatted runs."""
    doc = Document()
    doc.add_heading("Formatting Test", level=1)
    para = doc.add_paragraph()
    run = para.add_run("bold text")
    run.bold = True
    run = para.add_run(" and ")
    run = para.add_run("italic text")
    run.italic = True
    run = para.add_run(" with ")
    run = para.add_run("super")
    run.font.superscript = True

    p = tmp_path / "formatted.docx"
    doc.save(str(p))
    return p


@pytest.fixture
def tmp_tex(tmp_path: Path) -> Path:
    """Create a simple LaTeX project with root + input file."""
    methods = tmp_path / "methods.tex"
    methods.write_text(
        "\\section{Methods}\n"
        "\\label{sec:methods}\n"
        "\n"
        "We employ a 12-layer transformer.\n"
        "\n"
        "\\begin{equation}\n"
        "\\label{eq:loss}\n"
        "L = \\sum_{i=1}^{n} \\ell_i\n"
        "\\end{equation}\n"
        "\n"
        "Training uses AdamW optimizer.\n",
        encoding="utf-8",
    )

    root = tmp_path / "main.tex"
    root.write_text(
        "\\documentclass{article}\n"
        "\\begin{document}\n"
        "\n"
        "\\section{Introduction}\n"
        "This is the introduction paragraph.\n"
        "\n"
        "\\input{methods}\n"
        "\n"
        "\\end{document}\n",
        encoding="utf-8",
    )
    return root


@pytest.fixture
def empty_tex(tmp_path: Path) -> Path:
    """Create an empty LaTeX file."""
    p = tmp_path / "empty.tex"
    p.write_text(
        "\\documentclass{article}\n\\begin{document}\n\n\\end{document}\n",
        encoding="utf-8",
    )
    return p
