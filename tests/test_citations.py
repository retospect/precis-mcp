"""Tests for citation system — formatting, DOCX round-trip, orphan detection."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from precis.citations import (
    BIB_DEF_RE,
    BIB_ENTRY_STYLE,
    CITATION_REF_STYLE,
    CitationIndex,
    is_orphan_key,
    orphan_text,
    parse_ref_bookmark,
    ref_bookmark_name,
)
from precis.formatting import FormattedRun, markdown_to_runs, runs_to_markdown
from precis.parser.docx import DocxParser

# ---------------------------------------------------------------------------
# fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def cite_docx(tmp_path: Path) -> Path:
    """DOCX with a citation hyperlink, a bib entry, and a normal paragraph."""
    from precis.parser.docx import (
        _add_citation_hyperlink,
        _add_ref_bookmark,
        _ensure_styles,
    )

    doc = Document()
    _ensure_styles(doc)

    # Heading
    doc.add_heading("Introduction", level=1)

    # Paragraph with an inline citation hyperlink
    para = doc.add_paragraph("Results show improvement ")
    _add_citation_hyperlink(para, "[@smith2020]", "smith2020", doc)
    run = para.add_run(" in all conditions.")

    # Second citation of same key in another paragraph
    para2 = doc.add_paragraph("As noted by ")
    _add_citation_hyperlink(para2, "[@smith2020]", "smith2020", doc)
    run2 = para2.add_run(", the effect is robust.")

    # Bibliography heading
    doc.add_heading("References", level=1)

    # Bibliography entry with BibEntry style + ref_ bookmark
    bib_para = doc.add_paragraph()
    bib_para.style = doc.styles[BIB_ENTRY_STYLE]
    bib_para.add_run("Smith, J. (2020). Chicken in Shoes. J. Poultry, 42(1), 1-15.")
    _add_ref_bookmark(bib_para, "smith2020", doc)

    p = tmp_path / "cite.docx"
    doc.save(str(p))
    return p


@pytest.fixture
def orphan_docx(tmp_path: Path) -> Path:
    """DOCX with an orphaned CitationRef style (no hyperlink — simulates paste)."""
    from precis.parser.docx import _ensure_styles

    doc = Document()
    _ensure_styles(doc)

    doc.add_heading("Introduction", level=1)

    # Paragraph with a CitationRef-styled run but no hyperlink
    para = doc.add_paragraph("See ")
    run_el = para.add_run("(Smith 2020)")._element
    rpr = run_el.find(qn("w:rPr"))
    if rpr is None:
        rpr = etree.SubElement(run_el, qn("w:rPr"))
        run_el.insert(0, rpr)
    rstyle = etree.SubElement(rpr, qn("w:rStyle"))
    rstyle.set(qn("w:val"), CITATION_REF_STYLE)
    para.add_run(" for details.")

    p = tmp_path / "orphan.docx"
    doc.save(str(p))
    return p


# ---------------------------------------------------------------------------
# citations.py unit tests
# ---------------------------------------------------------------------------


class TestCitationHelpers:
    def test_ref_bookmark_name(self):
        assert ref_bookmark_name("smith2020") == "ref_smith2020"

    def test_parse_ref_bookmark(self):
        assert parse_ref_bookmark("ref_smith2020") == "smith2020"
        assert parse_ref_bookmark("ref_smith_jones_2020") == "smith_jones_2020"
        assert parse_ref_bookmark("cite_smith2020") is None
        assert parse_ref_bookmark("other") is None

    def test_orphan_helpers(self):
        assert is_orphan_key("?:some text")
        assert not is_orphan_key("smith2020")
        assert orphan_text("?:some text") == "some text"

    def test_bib_def_regex(self):
        m = BIB_DEF_RE.match("[@smith2020]: Smith, J. (2020). Title.")
        assert m is not None
        assert m.group(1) == "smith2020"
        assert m.group(2) == "Smith, J. (2020). Title."

        # Not a bib def (no colon after key)
        assert BIB_DEF_RE.match("[@smith2020] some text") is None


class TestCitationIndex:
    def test_register_and_render_author_year(self):
        idx = CitationIndex(style="author-year")
        idx.register_bib("smith2020", "Smith, J. (2020). Chicken in Shoes.")
        text, sup = idx.render_inline("smith2020")
        assert "Smith" in text
        assert "2020" in text
        assert not sup

    def test_numbered_style(self):
        idx = CitationIndex(style="numbered")
        idx.register_bib("smith2020", "Smith, J. (2020). Title.")
        idx.register_bib("jones2019", "Jones, A. (2019). Other.")
        t1, _ = idx.render_inline("smith2020")
        t2, _ = idx.render_inline("jones2019")
        assert t1 == "[1]"
        assert t2 == "[2]"
        # Second citation of smith should still be [1]
        t1b, _ = idx.render_inline("smith2020")
        assert t1b == "[1]"

    def test_superscript_style(self):
        idx = CitationIndex(style="superscript")
        idx.register_bib("smith2020", "Smith, J. (2020). Title.")
        text, sup = idx.render_inline("smith2020")
        assert text == "1"
        assert sup is True

    def test_unknown_key_author_year(self):
        idx = CitationIndex(style="author-year")
        text, sup = idx.render_inline("unknown_key")
        assert text == "[@unknown_key]"

    def test_round_trip_dict(self):
        idx = CitationIndex(style="numbered")
        idx.register_bib("smith2020", "Smith, J. (2020). Title.")
        idx.cite("smith2020")
        d = idx.to_dict()
        idx2 = CitationIndex.from_dict(d)
        assert idx2.style == "numbered"
        assert "smith2020" in idx2.entries
        assert idx2.cite_order == ["smith2020"]


# ---------------------------------------------------------------------------
# formatting.py citation tests
# ---------------------------------------------------------------------------


class TestFormattingCitations:
    def test_markdown_to_runs_citation(self):
        runs = markdown_to_runs("See [@smith2020] for details.")
        cite_runs = [r for r in runs if r.cite_key]
        assert len(cite_runs) == 1
        assert cite_runs[0].cite_key == "smith2020"

    def test_runs_to_markdown_citation(self):
        runs = [
            FormattedRun(text="See "),
            FormattedRun(text="(Smith 2020)", cite_key="smith2020"),
            FormattedRun(text=" for details."),
        ]
        md = runs_to_markdown(runs)
        assert "[@smith2020]" in md
        assert "(Smith 2020)" not in md  # cite_key overrides text

    def test_citation_round_trip_markdown(self):
        original = "Results from [@smith2020] and [@jones2019] show improvement."
        runs = markdown_to_runs(original)
        result = runs_to_markdown(runs)
        assert "[@smith2020]" in result
        assert "[@jones2019]" in result


# ---------------------------------------------------------------------------
# DOCX parser citation tests
# ---------------------------------------------------------------------------


class TestDocxCitationRead:
    def test_citation_hyperlink_detected(self, cite_docx: Path):
        parser = DocxParser()
        nodes = parser.parse(cite_docx)
        paras = [n for n in nodes if n.node_type == "p"]
        # First paragraph should contain [@smith2020]
        assert "[@smith2020]" in paras[0].text

    def test_multi_cite_same_key(self, cite_docx: Path):
        parser = DocxParser()
        nodes = parser.parse(cite_docx)
        paras = [n for n in nodes if n.node_type == "p"]
        # Both paragraphs should have [@smith2020]
        assert "[@smith2020]" in paras[0].text
        assert "[@smith2020]" in paras[1].text

    def test_bib_entry_detected(self, cite_docx: Path):
        parser = DocxParser()
        nodes = parser.parse(cite_docx)
        bibs = [n for n in nodes if n.node_type == "b"]
        assert len(bibs) == 1
        assert bibs[0].label == "smith2020"
        assert bibs[0].style == BIB_ENTRY_STYLE
        assert "Smith" in bibs[0].text

    def test_bib_entry_precis(self, cite_docx: Path):
        parser = DocxParser()
        nodes = parser.parse(cite_docx)
        bibs = [n for n in nodes if n.node_type == "b"]
        assert bibs[0].precis.startswith("smith2020:")

    def test_bib_entry_path(self, cite_docx: Path):
        parser = DocxParser()
        nodes = parser.parse(cite_docx)
        bibs = [n for n in nodes if n.node_type == "b"]
        # Should be under the References heading (S2.0.0.0)
        assert "b" in str(bibs[0].path)

    def test_orphaned_citation_detected(self, orphan_docx: Path):
        parser = DocxParser()
        nodes = parser.parse(orphan_docx)
        paras = [n for n in nodes if n.node_type == "p"]
        # Should contain orphaned citation marker
        assert any("[@?:" in p.text for p in paras)


class TestDocxCitationWrite:
    def test_append_bib_definition(self, tmp_path: Path):
        """Appending [@key]: text creates a BibEntry paragraph with ref_ bookmark."""
        doc = Document()
        doc.add_heading("References", level=1)
        p = tmp_path / "bib_write.docx"
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "[@smith2020]: Smith, J. (2020). Chicken in Shoes.")

        nodes = parser.parse(p)
        bibs = [n for n in nodes if n.node_type == "b"]
        assert len(bibs) == 1
        assert bibs[0].label == "smith2020"
        assert "Smith" in bibs[0].text
        # The [@key]: prefix should be stripped from displayed text
        assert "[@smith2020]" not in bibs[0].text

    def test_append_paragraph_with_citation(self, tmp_path: Path):
        """Appending text with [@key] creates a hyperlink in the DOCX."""
        doc = Document()
        doc.add_heading("Introduction", level=1)
        p = tmp_path / "cite_write.docx"
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "Results show improvement [@smith2020] overall.")

        nodes = parser.parse(p)
        paras = [n for n in nodes if n.node_type == "p"]
        assert len(paras) == 1
        assert "[@smith2020]" in paras[0].text

    def test_write_read_round_trip(self, tmp_path: Path):
        """Full round-trip: write bib + cite, read back, verify."""
        doc = Document()
        doc.add_heading("Introduction", level=1)
        p = tmp_path / "round_trip.docx"
        doc.save(str(p))

        parser = DocxParser()
        # Write a bib entry
        parser.append_node(p, "[@miller2023]: Miller, B. (2023). Quantum Poultry.")
        # Write a paragraph citing it
        parser.append_node(p, "As shown by [@miller2023], the effect is clear.")

        nodes = parser.parse(p)
        bibs = [n for n in nodes if n.node_type == "b"]
        paras = [n for n in nodes if n.node_type == "p"]

        assert len(bibs) == 1
        assert bibs[0].label == "miller2023"
        assert "[@miller2023]" in paras[0].text

    def test_multiple_bib_entries(self, tmp_path: Path):
        """Multiple bibliography entries get unique labels and paths."""
        doc = Document()
        doc.add_heading("References", level=1)
        p = tmp_path / "multi_bib.docx"
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "[@smith2020]: Smith, J. (2020). Title A.")
        parser.append_node(p, "[@jones2019]: Jones, A. (2019). Title B.")

        nodes = parser.parse(p)
        bibs = [n for n in nodes if n.node_type == "b"]
        assert len(bibs) == 2
        labels = {b.label for b in bibs}
        assert labels == {"smith2020", "jones2019"}
        # Paths should be different
        assert str(bibs[0].path) != str(bibs[1].path)

    def test_multiple_cites_same_key(self, tmp_path: Path):
        """Multiple citations of the same key all round-trip."""
        doc = Document()
        doc.add_heading("Test", level=1)
        p = tmp_path / "multi_cite.docx"
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "First [@smith2020] and second [@smith2020] citations.")

        nodes = parser.parse(p)
        paras = [n for n in nodes if n.node_type == "p"]
        assert paras[0].text.count("[@smith2020]") == 2


# ---------------------------------------------------------------------------
# Deep round-trip tests — XML structure + multi-step workflows
# ---------------------------------------------------------------------------


class TestCitationRoundTrip:
    def test_full_document_lifecycle(self, tmp_path: Path):
        """Build a complete document with citations and bib, verify everything."""
        p = tmp_path / "lifecycle.docx"
        doc = Document()
        doc.save(str(p))

        parser = DocxParser()

        # Step 1: Add structure
        parser.append_node(p, "# Introduction")
        parser.append_node(p, "The field has advanced rapidly [@smith2020].")
        parser.append_node(
            p, "Recent work [@jones2019] confirms earlier findings [@smith2020]."
        )
        parser.append_node(p, "# References")
        parser.append_node(
            p, "[@smith2020]: Smith, J. (2020). Chicken in Shoes. J. Poultry."
        )
        parser.append_node(
            p, "[@jones2019]: Jones, A. (2019). Quantum Poultry Dynamics."
        )

        # Step 2: Parse and verify
        nodes = parser.parse(p)
        headings = [n for n in nodes if n.node_type == "h"]
        paras = [n for n in nodes if n.node_type == "p"]
        bibs = [n for n in nodes if n.node_type == "b"]

        assert len(headings) == 2
        assert len(paras) == 2
        assert len(bibs) == 2

        # Verify inline citations survived
        assert "[@smith2020]" in paras[0].text
        assert "[@jones2019]" in paras[1].text
        assert paras[1].text.count("[@smith2020]") == 1
        assert paras[1].text.count("[@jones2019]") == 1

        # Verify bib entries
        labels = {b.label for b in bibs}
        assert labels == {"smith2020", "jones2019"}
        assert all(b.node_type == "b" for b in bibs)
        assert all("b" in str(b.path) for b in bibs)

    def test_xml_hyperlink_structure(self, tmp_path: Path):
        """Verify the actual XML has w:hyperlink with correct w:anchor."""
        p = tmp_path / "xml_check.docx"
        doc = Document()
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "See [@miller2023] for details.")

        # Re-open and inspect raw XML
        doc2 = Document(str(p))
        body = doc2.element.body
        hyperlinks = list(body.iter(qn("w:hyperlink")))
        assert len(hyperlinks) == 1
        assert hyperlinks[0].get(qn("w:anchor")) == "ref_miller2023"

        # Verify CitationRef style on the run inside
        runs_in_link = hyperlinks[0].findall(qn("w:r"))
        assert len(runs_in_link) >= 1
        rpr = runs_in_link[0].find(qn("w:rPr"))
        assert rpr is not None
        rstyle = rpr.find(qn("w:rStyle"))
        assert rstyle is not None
        assert rstyle.get(qn("w:val")) == CITATION_REF_STYLE

    def test_xml_bookmark_structure(self, tmp_path: Path):
        """Verify bib entry has ref_ bookmark in the XML."""
        p = tmp_path / "bm_check.docx"
        doc = Document()
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "[@lee2021]: Lee, K. (2021). Deep Poultry Networks.")

        doc2 = Document(str(p))
        body = doc2.element.body
        bookmarks = list(body.iter(qn("w:bookmarkStart")))
        ref_bookmarks = [
            bm for bm in bookmarks if bm.get(qn("w:name"), "").startswith("ref_")
        ]
        assert len(ref_bookmarks) == 1
        assert ref_bookmarks[0].get(qn("w:name")) == "ref_lee2021"

        # Verify matching bookmarkEnd exists
        bm_id = ref_bookmarks[0].get(qn("w:id"))
        ends = [e for e in body.iter(qn("w:bookmarkEnd")) if e.get(qn("w:id")) == bm_id]
        assert len(ends) == 1

    def test_xml_bib_entry_style(self, tmp_path: Path):
        """Verify bib entry paragraph has BibEntry style in XML."""
        p = tmp_path / "style_check.docx"
        doc = Document()
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "[@wang2022]: Wang, X. (2022). Title.")

        doc2 = Document(str(p))
        # Find paragraph with BibEntry style
        for para in doc2.paragraphs:
            if para.style and para.style.name == BIB_ENTRY_STYLE:
                assert "Wang" in para.text
                break
        else:
            pytest.fail("No paragraph with BibEntry style found")

    def test_edit_paragraph_preserves_citations(self, tmp_path: Path):
        """Edit a paragraph via write_node — citations should survive if in new text."""
        p = tmp_path / "edit_cite.docx"
        doc = Document()
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "# Intro")
        parser.append_node(p, "Old text with [@smith2020] citation.")

        nodes = parser.parse(p)
        para = [n for n in nodes if n.node_type == "p"][0]

        # Edit the paragraph, keeping the citation
        parser.write_node(p, para, "New text still cites [@smith2020] here.")

        nodes2 = parser.parse(p)
        paras2 = [n for n in nodes2 if n.node_type == "p"]
        assert len(paras2) == 1
        assert "[@smith2020]" in paras2[0].text
        assert "New text" in paras2[0].text

    def test_mixed_formatting_and_citations(self, tmp_path: Path):
        """Bold/italic text mixed with citations round-trips correctly."""
        p = tmp_path / "mixed.docx"
        doc = Document()
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "# Results")
        parser.append_node(
            p, "**Significant** results from [@smith2020] with *p < 0.05*."
        )

        nodes = parser.parse(p)
        paras = [n for n in nodes if n.node_type == "p"]
        text = paras[0].text
        assert "[@smith2020]" in text
        assert "**Significant**" in text
        assert "*p < 0.05*" in text

    def test_insert_after_with_citation(self, tmp_path: Path):
        """insert_after with citation text creates proper hyperlink."""
        p = tmp_path / "insert.docx"
        doc = Document()
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "# Intro")
        parser.append_node(p, "First paragraph.")

        nodes = parser.parse(p)
        anchor = [n for n in nodes if n.node_type == "p"][0]

        parser.insert_after(p, anchor, "Added with [@ref2024] inline.")

        nodes2 = parser.parse(p)
        paras2 = [n for n in nodes2 if n.node_type == "p"]
        assert len(paras2) == 2
        assert "[@ref2024]" in paras2[1].text

    def test_insert_before_with_citation(self, tmp_path: Path):
        """insert_before with citation text creates proper hyperlink."""
        p = tmp_path / "insert_b.docx"
        doc = Document()
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "# Intro")
        parser.append_node(p, "Second paragraph.")

        nodes = parser.parse(p)
        anchor = [n for n in nodes if n.node_type == "p"][0]

        parser.insert_before(p, anchor, "Prepended with [@pre2024] citation.")

        nodes2 = parser.parse(p)
        paras2 = [n for n in nodes2 if n.node_type == "p"]
        assert len(paras2) == 2
        assert "[@pre2024]" in paras2[0].text

    def test_double_round_trip(self, tmp_path: Path):
        """Write → read → write modified → read — verify integrity."""
        p = tmp_path / "double.docx"
        doc = Document()
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "# Paper")
        parser.append_node(p, "Claim supported by [@a2020] and [@b2021].")
        parser.append_node(p, "[@a2020]: Alpha, A. (2020). First.")
        parser.append_node(p, "[@b2021]: Beta, B. (2021). Second.")

        # First read
        nodes1 = parser.parse(p)
        para1 = [n for n in nodes1 if n.node_type == "p"][0]
        bibs1 = [n for n in nodes1 if n.node_type == "b"]
        assert "[@a2020]" in para1.text
        assert "[@b2021]" in para1.text
        assert len(bibs1) == 2

        # Modify the paragraph — add a third citation
        parser.write_node(
            p, para1, "Claim supported by [@a2020], [@b2021], and [@c2022]."
        )

        # Second read
        nodes2 = parser.parse(p)
        para2 = [n for n in nodes2 if n.node_type == "p"][0]
        bibs2 = [n for n in nodes2 if n.node_type == "b"]

        assert "[@a2020]" in para2.text
        assert "[@b2021]" in para2.text
        assert "[@c2022]" in para2.text
        # Bibs unchanged
        assert len(bibs2) == 2

    def test_bib_text_no_prefix_in_docx(self, tmp_path: Path):
        """The [@key]: prefix is metadata — it should NOT appear in the DOCX text."""
        p = tmp_path / "no_prefix.docx"
        doc = Document()
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "[@zz2020]: Zeta, Z. (2020). Title Here.")

        # Check raw DOCX paragraph text
        doc2 = Document(str(p))
        for para in doc2.paragraphs:
            if "Zeta" in para.text:
                assert "[@zz2020]" not in para.text
                assert para.text.strip().startswith("Zeta")
                break
        else:
            pytest.fail("Bib entry paragraph not found")

    def test_citation_in_tracked_change(self, tmp_path: Path):
        """Track-changes replacement with citation text."""
        p = tmp_path / "tracked.docx"
        doc = Document()
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "# Intro")
        parser.append_node(p, "Old content here.")

        nodes = parser.parse(p)
        para = [n for n in nodes if n.node_type == "p"][0]

        # Apply tracked change with a citation
        parser.write_tracked(p, para, "New content with [@ref2025] cited.")

        # File should still be valid
        nodes2 = parser.parse(p)
        assert len(nodes2) >= 1  # at least heading survives

    def test_empty_citation_key_ignored(self, tmp_path: Path):
        """A malformed [@] should not crash — treated as plain text."""
        p = tmp_path / "empty_key.docx"
        doc = Document()
        doc.save(str(p))

        parser = DocxParser()
        parser.append_node(p, "# Test")
        parser.append_node(p, "This has [@] which is not a valid citation.")

        nodes = parser.parse(p)
        paras = [n for n in nodes if n.node_type == "p"]
        assert len(paras) == 1
        # Should not crash, text preserved
        assert "[@]" in paras[0].text or "[@" in paras[0].text
