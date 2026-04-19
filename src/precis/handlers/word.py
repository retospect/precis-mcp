"""DOCX handler — full read/write with tracked changes and comments.

Integrates DocxParser from precis v1 with the v2 handler protocol.
Requires the ``word`` extra: ``pip install precis[word]``.
"""

from __future__ import annotations

import os
import tempfile
from datetime import UTC
from pathlib import Path

from docx import Document
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.ns import qn
from lxml import etree

from precis.citations import (
    BIB_DEF_RE,
    BIB_ENTRY_STYLE,
    CITATION_REF_STYLE,
    ORPHAN_PREFIX,
    parse_ref_bookmark,
    ref_bookmark_name,
)
from precis.formatting import (
    LIST_ITEM_RE,
    FormattedRun,
    list_prefix,
    markdown_to_runs,
    parse_list_prefix,
    runs_to_markdown,
)
from precis.handlers._file_base import FileHandlerBase
from precis.protocol import Node, PathCounter, PrecisError, make_slug, resolve_slug

# Hardened XML parser for reading .docx XML parts.  Disables external
# entity resolution, DTD loading, and network access so a crafted .docx
# cannot trigger XXE, billion-laughs, or SSRF attacks even if the file
# comes from an untrusted source.  Applies only to parsing; element
# creation (etree.Element / etree.SubElement) is intrinsically safe.
_SAFE_XML_PARSER = etree.XMLParser(
    resolve_entities=False,
    no_network=True,
    dtd_validation=False,
    load_dtd=False,
    huge_tree=False,
)


def _parse_xml(blob: bytes):
    """Parse a DOCX-internal XML blob with the hardened parser."""
    return etree.fromstring(blob, parser=_SAFE_XML_PARSER)


# ─── DOCX list constants ──────────────────────────────────────────

_RT_NUMBERING = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering"
)

_LIST_BULLET_STYLES = ["List Bullet", "List Bullet 2", "List Bullet 3"]
_LIST_NUMBER_STYLES = ["List Number", "List Number 2", "List Number 3"]


class WordHandler(FileHandlerBase):
    """Handler for .docx files."""

    extensions = {".docx"}

    # ── Parser implementation ───────────────────────────────────────

    def parse(self, path: Path) -> list[Node]:
        """Parse a DOCX file into a list of Nodes."""
        doc = Document(str(path))
        comments_map = _parse_comments(doc)
        num_map = _build_numbering_map(doc)
        counter = PathCounter()
        nodes: list[Node] = []
        slug_counts: dict[str, int] = {}

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break
                if para is None:
                    continue

                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else "Normal"

                heading_level = _heading_level(style_name)
                if heading_level:
                    node_path = counter.next_heading(heading_level)
                    base_slug = make_slug(text)
                    slug = resolve_slug(base_slug, slug_counts)
                    nodes.append(
                        Node(
                            slug=slug,
                            path=node_path,
                            node_type="h",
                            text=text,
                            precis=text,
                            style=style_name,
                            comments=_collect_para_comments(element, comments_map),
                        )
                    )
                elif style_name == BIB_ENTRY_STYLE:
                    node_path = counter.next_child("b")
                    md_text = _para_to_markdown(para)
                    label = _extract_bib_label(element)
                    base_slug = make_slug(md_text)
                    slug = resolve_slug(base_slug, slug_counts)
                    nodes.append(
                        Node(
                            slug=slug,
                            path=node_path,
                            node_type="b",
                            text=md_text,
                            precis=_bib_precis(label, text),
                            style=style_name,
                            label=label,
                            comments=_collect_para_comments(element, comments_map),
                        )
                    )
                else:
                    node_path = counter.next_child("p")
                    md_text = _para_to_markdown(para)
                    li = _get_list_info(element, style_name, num_map)
                    if li:
                        md_text = list_prefix(li[0], li[1]) + md_text
                    base_slug = make_slug(md_text)
                    slug = resolve_slug(base_slug, slug_counts)
                    nodes.append(
                        Node(
                            slug=slug,
                            path=node_path,
                            node_type="p",
                            text=md_text,
                            style=style_name,
                            comments=_collect_para_comments(element, comments_map),
                        )
                    )

            elif tag == "tbl":
                table = None
                for t in doc.tables:
                    if t._element is element:
                        table = t
                        break
                if table is None:
                    continue

                md_table = _table_to_markdown(table)
                synopsis = _table_synopsis(table)
                node_path = counter.next_child("t")
                base_slug = make_slug(md_table)
                slug = resolve_slug(base_slug, slug_counts)
                nodes.append(
                    Node(
                        slug=slug,
                        path=node_path,
                        node_type="t",
                        text=md_table,
                        precis=synopsis,
                        style="Table",
                    )
                )

        return nodes

    def source_files(self, path: Path) -> list[Path]:
        return [path]

    def write_node(self, path: Path, node: Node, new_text: str) -> None:
        """Replace a node's text content."""
        doc = Document(str(path))
        element = self._find_element(doc, node)
        if element is None:
            raise PrecisError(f"Node not found in document: {node.slug}")

        if node.node_type in ("p", "h"):
            para = _element_to_para(doc, element)
            if para is None:
                raise PrecisError(f"Paragraph not found: {node.slug}")
            _set_para_text(para, new_text, node.node_type == "h")

        _atomic_save(doc, path)

    def insert_after(
        self, path: Path, anchor: Node, new_text: str, heading_level: int = 0
    ) -> None:
        doc = Document(str(path))
        element = self._find_element(doc, anchor)
        if element is None:
            raise PrecisError(f"Anchor not found: {anchor.slug}")

        new_para = _make_paragraph(doc, new_text, heading_level)
        element.addnext(new_para._element)
        _atomic_save(doc, path)

    def insert_before(
        self, path: Path, anchor: Node, new_text: str, heading_level: int = 0
    ) -> None:
        doc = Document(str(path))
        element = self._find_element(doc, anchor)
        if element is None:
            raise PrecisError(f"Anchor not found: {anchor.slug}")

        new_para = _make_paragraph(doc, new_text, heading_level)
        element.addprevious(new_para._element)
        _atomic_save(doc, path)

    def delete_node(self, path: Path, node: Node) -> None:
        doc = Document(str(path))
        element = self._find_element(doc, node)
        if element is None:
            raise PrecisError(f"Node not found: {node.slug}")
        element.getparent().remove(element)
        _atomic_save(doc, path)

    def append_node(self, path: Path, new_text: str, heading_level: int = 0) -> None:
        doc = Document(str(path))
        if "\n" in new_text and LIST_ITEM_RE.match(new_text.strip()):
            for line in new_text.split("\n"):
                line = line.strip()
                if line:
                    para = _make_paragraph(doc, line)
                    doc.element.body.append(para._element)
        else:
            new_para = _make_paragraph(doc, new_text, heading_level)
            doc.element.body.append(new_para._element)
        _atomic_save(doc, path)

    def move_nodes(self, path: Path, nodes: list[Node], after: Node) -> None:
        doc = Document(str(path))
        after_el = self._find_element(doc, after)
        if after_el is None:
            raise PrecisError(f"Target not found: {after.slug}")

        elements = []
        for n in nodes:
            el = self._find_element(doc, n)
            if el is not None:
                elements.append(el)

        for el in elements:
            el.getparent().remove(el)

        insert_point = after_el
        for el in elements:
            insert_point.addnext(el)
            insert_point = el

        _atomic_save(doc, path)

    def write_tracked(
        self, path: Path, node: Node, new_text: str, author: str = "precis"
    ) -> None:
        """Replace node text with track-changes markup."""
        doc = Document(str(path))
        element = self._find_element(doc, node)
        if element is None:
            raise PrecisError(f"Node not found: {node.slug}")

        if node.node_type in ("p", "h"):
            para = _element_to_para(doc, element)
            if para is None:
                raise PrecisError(f"Paragraph not found: {node.slug}")
            _inject_tracked_replace(para, new_text, author)

        _atomic_save(doc, path)

    def write_comment(
        self, path: Path, node: Node, text: str, author: str = "precis"
    ) -> int:
        """Add a Word comment (margin annotation). Returns comment ID."""
        doc = Document(str(path))
        element = self._find_element(doc, node)
        if element is None:
            raise PrecisError(f"Node not found: {node.slug}")

        comment_id = _inject_comment(doc, element, text, author)
        _atomic_save(doc, path)
        return comment_id

    # ── Element finding ─────────────────────────────────────────────

    def _find_element(self, doc: Document, node: Node):
        """Find the XML element matching a node by reparsing and matching."""
        counter = PathCounter()
        slug_counts: dict[str, int] = {}
        num_map = _build_numbering_map(doc)

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                para = _element_to_para(doc, element)
                if para is None or not para.text.strip():
                    continue
                style_name = para.style.name if para.style else "Normal"
                heading_level = _heading_level(style_name)
                if heading_level:
                    counter.next_heading(heading_level)
                    text = para.text.strip()
                elif style_name == BIB_ENTRY_STYLE:
                    counter.next_child("b")
                    text = _para_to_markdown(para)
                else:
                    counter.next_child("p")
                    text = _para_to_markdown(para)
                    li = _get_list_info(element, style_name, num_map)
                    if li:
                        text = list_prefix(li[0], li[1]) + text
                base_slug = make_slug(text)
                slug = resolve_slug(base_slug, slug_counts)
                if slug == node.slug:
                    return element

            elif tag == "tbl":
                table = _element_to_table(doc, element)
                if table is None:
                    continue
                counter.next_child("t")
                md = _table_to_markdown(table)
                base_slug = make_slug(md)
                slug = resolve_slug(base_slug, slug_counts)
                if slug == node.slug:
                    return element

        return None


# ─── Module-level helpers (ported from precis-mcp DocxParser) ───────


def _element_to_para(doc, element):
    for p in doc.paragraphs:
        if p._element is element:
            return p
    return None


def _element_to_table(doc, element):
    for t in doc.tables:
        if t._element is element:
            return t
    return None


def _atomic_save(doc: Document, path: Path) -> None:
    fd, tmp = tempfile.mkstemp(dir=path.parent, suffix=".docx.tmp")
    os.close(fd)
    try:
        doc.save(tmp)
        os.replace(tmp, path)
    except Exception:
        try:
            os.unlink(tmp)
        except OSError:
            pass
        raise


def _heading_level(style_name: str) -> int:
    """Extract heading level from style name, or 0.

    Title → 1, Heading 1 → 1, Heading 2 → 2, Heading 3 → 3, Heading 4 → 4.
    """
    if not style_name:
        return 0
    s = style_name.lower()
    if s == "title":
        return 1
    if s.startswith("heading"):
        rest = s[7:].strip()
        if rest.isdigit():
            level = int(rest)
            if 1 <= level <= 4:
                return level
    return 0


def _build_numbering_map(doc) -> dict[tuple[str, int], str]:
    """Map (numId, ilvl) → 'bullet' | 'number' from word/numbering.xml."""
    result: dict[tuple[str, int], str] = {}
    try:
        npart = None
        for rel in doc.part.rels.values():
            if rel.reltype == _RT_NUMBERING:
                npart = rel.target_part
                break
        if npart is None:
            return result

        root = _parse_xml(npart.blob)

        # abstractNumId → {ilvl: format}
        abs_map: dict[str, dict[int, str]] = {}
        for abs_el in root.findall(qn("w:abstractNum")):
            abs_id = abs_el.get(qn("w:abstractNumId"), "")
            levels: dict[int, str] = {}
            for lvl in abs_el.findall(qn("w:lvl")):
                ilvl = int(lvl.get(qn("w:ilvl"), "0"))
                fmt_el = lvl.find(qn("w:numFmt"))
                fmt = fmt_el.get(qn("w:val"), "") if fmt_el is not None else ""
                levels[ilvl] = "bullet" if fmt == "bullet" else "number"
            abs_map[abs_id] = levels

        # numId → expand via abstractNumId
        for num_el in root.findall(qn("w:num")):
            nid = num_el.get(qn("w:numId"), "")
            aid_el = num_el.find(qn("w:abstractNumId"))
            if aid_el is not None:
                aid = aid_el.get(qn("w:val"), "")
                for ilvl, fmt in abs_map.get(aid, {}).items():
                    result[(nid, ilvl)] = fmt
    except Exception:
        pass
    return result


def _get_list_info(
    element, style_name: str, num_map: dict[tuple[str, int], str]
) -> tuple[str, int] | None:
    """Detect list paragraph. Returns ('bullet'|'number', ilvl) or None."""
    ilvl = 0
    has_numPr = False

    pPr = element.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            nid_el = numPr.find(qn("w:numId"))
            if nid_el is not None:
                nid = nid_el.get(qn("w:val"), "0")
                if nid != "0":
                    has_numPr = True
                    ilvl_el = numPr.find(qn("w:ilvl"))
                    ilvl = (
                        int(ilvl_el.get(qn("w:val"), "0")) if ilvl_el is not None else 0
                    )

                    # Try numbering definition first
                    fmt = num_map.get((nid, ilvl))
                    if fmt:
                        return (fmt, ilvl)

    # Fallback: check style name (covers inherited numPr from style def)
    sn = style_name.lower()
    if "bullet" in sn or "list bullet" in sn:
        return ("bullet", ilvl)
    if "number" in sn or "list number" in sn:
        return ("number", ilvl)
    if sn == "list paragraph" and has_numPr:
        return ("bullet", ilvl)

    if has_numPr:
        return ("bullet", ilvl)

    return None


def _local_tag(el) -> str:
    tag = el.tag
    if "}" in tag:
        return tag.split("}")[-1]
    return tag


def _run_element_text(r_el) -> str:
    parts = []
    for t_el in r_el.findall(qn("w:t")):
        if t_el.text:
            parts.append(t_el.text)
    return "".join(parts)


def _has_char_style(r_el, style_name: str) -> bool:
    rpr = r_el.find(qn("w:rPr"))
    if rpr is None:
        return False
    rstyle = rpr.find(qn("w:rStyle"))
    if rstyle is None:
        return False
    return rstyle.get(qn("w:val")) == style_name


def _run_element_to_formatted(r_el, text: str) -> FormattedRun:
    rpr = r_el.find(qn("w:rPr"))
    bold = False
    italic = False
    superscript = False
    subscript = False
    strike = False

    if rpr is not None:
        bold = rpr.find(qn("w:b")) is not None
        italic = rpr.find(qn("w:i")) is not None
        strike = rpr.find(qn("w:strike")) is not None
        vert = rpr.find(qn("w:vertAlign"))
        if vert is not None:
            val = vert.get(qn("w:val"), "")
            superscript = val == "superscript"
            subscript = val == "subscript"

    return FormattedRun(
        text=text,
        bold=bold,
        italic=italic,
        superscript=superscript,
        subscript=subscript,
        strike=strike,
    )


def _para_to_markdown(para) -> str:
    """Convert a python-docx paragraph to Markdown text, including citations."""
    runs_list: list[FormattedRun] = []

    for child in para._element:
        tag = _local_tag(child)

        if tag == "hyperlink":
            anchor = child.get(qn("w:anchor"), "")
            cite_key = ""
            if anchor.startswith("ref_"):
                cite_key = anchor[4:]

            for r_el in child.findall(qn("w:r")):
                text = _run_element_text(r_el)
                if not text:
                    continue
                if cite_key:
                    runs_list.append(FormattedRun(text=text, cite_key=cite_key))
                else:
                    runs_list.append(_run_element_to_formatted(r_el, text))

        elif tag == "r":
            text = _run_element_text(child)
            if not text:
                continue
            if _has_char_style(child, CITATION_REF_STYLE):
                orphan_key = f"{ORPHAN_PREFIX}{text}"
                runs_list.append(FormattedRun(text=text, cite_key=orphan_key))
            else:
                runs_list.append(_run_element_to_formatted(child, text))

    return runs_to_markdown(runs_list)


def _set_para_text(
    para, markdown_text: str, is_heading: bool = False, doc=None
) -> None:
    for child in list(para._element):
        tag = _local_tag(child)
        if tag in ("r", "hyperlink"):
            para._element.remove(child)

    if is_heading:
        text = markdown_text.lstrip("#").strip()
        run = para.add_run(text)
        return

    formatted_runs = markdown_to_runs(markdown_text)
    _write_formatted_runs(para, formatted_runs, doc)


def _write_formatted_runs(para, formatted_runs: list[FormattedRun], doc=None) -> None:
    for fr in formatted_runs:
        if fr.cite_key and doc is not None:
            _add_citation_hyperlink(para, fr.text, fr.cite_key, doc)
        else:
            run = para.add_run(fr.text)
            if fr.bold:
                run.bold = True
            if fr.italic:
                run.italic = True
            if fr.superscript:
                run.font.superscript = True
            if fr.subscript:
                run.font.subscript = True
            if fr.strike:
                run.font.strike = True


def _make_paragraph(doc, text: str, heading_level: int = 0):
    bib_m = BIB_DEF_RE.match(text)
    if bib_m:
        key = bib_m.group(1)
        ref_text = bib_m.group(2)
        _ensure_styles(doc)
        para = doc.add_paragraph()
        para.style = doc.styles[BIB_ENTRY_STYLE]
        formatted_runs = markdown_to_runs(ref_text)
        _write_formatted_runs(para, formatted_runs, doc)
        _add_ref_bookmark(para, key, doc)
        return para

    li = parse_list_prefix(text)
    if li:
        list_type, indent_level, content = li
        return _make_list_paragraph(doc, content, list_type, indent_level)

    if text.startswith("#"):
        level = 0
        stripped = text
        while stripped.startswith("#"):
            level += 1
            stripped = stripped[1:]
        stripped = stripped.strip()
        if level >= 1:
            heading_level = min(level, 4)
            text = stripped

    if heading_level:
        para = doc.add_paragraph(text)
        # # → Title, ## → Heading 1, ### → Heading 2, #### → Heading 3
        if heading_level == 1:
            para.style = doc.styles["Title"]
        else:
            para.style = doc.styles[f"Heading {heading_level - 1}"]
    else:
        para = doc.add_paragraph()
        formatted_runs = markdown_to_runs(text)
        _write_formatted_runs(para, formatted_runs, doc)

    return para


def _make_list_paragraph(doc, text: str, list_type: str, indent_level: int = 0):
    """Create a paragraph with a Word list style."""
    if list_type == "number":
        styles = _LIST_NUMBER_STYLES
    else:
        styles = _LIST_BULLET_STYLES
    style_name = styles[min(indent_level, len(styles) - 1)]

    para = doc.add_paragraph()
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        pass
    formatted_runs = markdown_to_runs(text)
    _write_formatted_runs(para, formatted_runs, doc)
    return para


def _table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")

    if len(rows) >= 1:
        ncols = len(table.rows[0].cells)
        sep = "| " + " | ".join(["---"] * ncols) + " |"
        rows.insert(1, sep)

    return "\n".join(rows)


def _table_synopsis(table) -> str:
    if not table.rows:
        return "[empty table]"
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    nrows = len(table.rows) - 1
    synopsis = "|".join(headers) + f" ({nrows}r)"
    return synopsis


def _extract_bib_label(p_el) -> str:
    for bm in p_el.findall(qn("w:bookmarkStart")):
        name = bm.get(qn("w:name"), "")
        key = parse_ref_bookmark(name)
        if key:
            return key
    return ""


def _bib_precis(label: str, text: str) -> str:
    short = text[:60].rstrip()
    if len(text) > 60:
        short += "…"
    if label:
        return f"{label}: {short}"
    return short


# ─── Track changes ──────────────────────────────────────────────────


def _inject_tracked_replace(para, new_text: str, author: str) -> None:
    from datetime import datetime

    now = datetime.now(UTC).strftime("%Y-%m-%dT%H:%M:%SZ")

    for run_el in list(para._element.findall(qn("w:r"))):
        del_el = etree.SubElement(para._element, qn("w:del"))
        del_el.set(qn("w:author"), author)
        del_el.set(qn("w:date"), now)
        para._element.remove(run_el)
        for t_el in run_el.findall(qn("w:t")):
            t_el.tag = qn("w:delText")
            t_el.set(qn("xml:space"), "preserve")
        del_el.append(run_el)

    ins_el = etree.SubElement(para._element, qn("w:ins"))
    ins_el.set(qn("w:author"), author)
    ins_el.set(qn("w:date"), now)

    formatted_runs = markdown_to_runs(new_text)
    for fr in formatted_runs:
        run_el = etree.SubElement(ins_el, qn("w:r"))
        t_el = etree.SubElement(run_el, qn("w:t"))
        t_el.text = fr.text
        t_el.set(qn("xml:space"), "preserve")

        if fr.bold or fr.italic or fr.superscript or fr.subscript or fr.strike:
            rpr = etree.SubElement(run_el, qn("w:rPr"))
            run_el.insert(0, rpr)
            if fr.bold:
                etree.SubElement(rpr, qn("w:b"))
            if fr.italic:
                etree.SubElement(rpr, qn("w:i"))
            if fr.strike:
                etree.SubElement(rpr, qn("w:strike"))
            if fr.superscript:
                vertAlign = etree.SubElement(rpr, qn("w:vertAlign"))
                vertAlign.set(qn("w:val"), "superscript")
            if fr.subscript:
                vertAlign = etree.SubElement(rpr, qn("w:vertAlign"))
                vertAlign.set(qn("w:val"), "subscript")


# ─── Word comments ──────────────────────────────────────────────────

_COMMENTS_URI = PackURI("/word/comments.xml")
_RT_COMMENTS = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
_COMMENTS_CT = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)


def _get_comments_part(doc):
    for rel in doc.part.rels.values():
        if rel.reltype == _RT_COMMENTS:
            return rel.target_part
    return None


def _get_or_create_comments_part(doc):
    part = _get_comments_part(doc)
    if part is not None:
        return part
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    root = etree.Element(qn("w:comments"), nsmap=nsmap)
    blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    part = Part(_COMMENTS_URI, _COMMENTS_CT, blob, doc.part.package)
    doc.part.relate_to(part, _RT_COMMENTS)
    return part


def _parse_comments(doc) -> dict[int, dict]:
    part = _get_comments_part(doc)
    if part is None:
        return {}
    root = _parse_xml(part.blob)
    result = {}
    for comment_el in root.findall(qn("w:comment")):
        try:
            cid = int(comment_el.get(qn("w:id"), "0"))
        except ValueError:
            continue
        author = comment_el.get(qn("w:author"), "")
        texts = []
        for p_el in comment_el.findall(qn("w:p")):
            for r_el in p_el.findall(qn("w:r")):
                for t_el in r_el.findall(qn("w:t")):
                    if t_el.text:
                        texts.append(t_el.text)
        result[cid] = {"author": author, "text": " ".join(texts)}
    return result


def _next_comment_id(doc) -> int:
    part = _get_comments_part(doc)
    if part is None:
        return 1
    root = _parse_xml(part.blob)
    max_id = 0
    for comment_el in root.findall(qn("w:comment")):
        try:
            cid = int(comment_el.get(qn("w:id"), "0"))
            if cid > max_id:
                max_id = cid
        except ValueError:
            pass
    return max_id + 1


def _collect_para_comments(element, comments_map: dict) -> list[dict]:
    comments = []
    seen_ids: set[int] = set()
    for el in element.findall(qn("w:commentRangeStart")):
        try:
            cid = int(el.get(qn("w:id"), "0"))
        except ValueError:
            continue
        if cid in comments_map and cid not in seen_ids:
            seen_ids.add(cid)
            cm = comments_map[cid]
            comments.append({"id": cid, "author": cm["author"], "text": cm["text"]})
    for r_el in element.findall(qn("w:r")):
        for ref_el in r_el.findall(qn("w:commentReference")):
            try:
                cid = int(ref_el.get(qn("w:id"), "0"))
            except ValueError:
                continue
            if cid in comments_map and cid not in seen_ids:
                seen_ids.add(cid)
                cm = comments_map[cid]
                comments.append({"id": cid, "author": cm["author"], "text": cm["text"]})
    return comments


def _inject_comment(doc, para_element, text: str, author: str = "precis") -> int:
    from datetime import datetime

    comment_id = _next_comment_id(doc)
    now = datetime.now(UTC).strftime("%Y-%m-%dT%H:%M:%SZ")
    cid_str = str(comment_id)

    comments_part = _get_or_create_comments_part(doc)

    if hasattr(comments_part, "_element") and comments_part._element is not None:
        root = comments_part._element
    else:
        root = _parse_xml(comments_part.blob)

    comment_el = etree.SubElement(root, qn("w:comment"))
    comment_el.set(qn("w:id"), cid_str)
    comment_el.set(qn("w:author"), author)
    comment_el.set(qn("w:date"), now)
    comment_el.set(qn("w:initials"), author[:2].upper())

    p_el = etree.SubElement(comment_el, qn("w:p"))
    r_el = etree.SubElement(p_el, qn("w:r"))
    t_el = etree.SubElement(r_el, qn("w:t"))
    t_el.text = text
    t_el.set(qn("xml:space"), "preserve")

    if not hasattr(comments_part, "_element") or comments_part._element is None:
        comments_part._blob = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    range_start = etree.Element(qn("w:commentRangeStart"))
    range_start.set(qn("w:id"), cid_str)

    ppr = para_element.find(qn("w:pPr"))
    first_run = para_element.find(qn("w:r"))
    if first_run is not None:
        first_run.addprevious(range_start)
    elif ppr is not None:
        ppr.addnext(range_start)
    else:
        para_element.insert(0, range_start)

    range_end = etree.SubElement(para_element, qn("w:commentRangeEnd"))
    range_end.set(qn("w:id"), cid_str)

    ref_run = etree.SubElement(para_element, qn("w:r"))
    ref_el = etree.SubElement(ref_run, qn("w:commentReference"))
    ref_el.set(qn("w:id"), cid_str)

    return comment_id


# ─── Style and citation helpers ─────────────────────────────────────


def _ensure_styles(doc) -> None:
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Inches

    styles = doc.styles

    try:
        styles[CITATION_REF_STYLE]
    except KeyError:
        styles.add_style(CITATION_REF_STYLE, WD_STYLE_TYPE.CHARACTER)

    try:
        styles[BIB_ENTRY_STYLE]
    except KeyError:
        style = styles.add_style(BIB_ENTRY_STYLE, WD_STYLE_TYPE.PARAGRAPH)
        style.paragraph_format.first_line_indent = Inches(-0.5)
        style.paragraph_format.left_indent = Inches(0.5)


def _add_citation_hyperlink(para, text: str, cite_key: str, doc) -> None:
    _ensure_styles(doc)
    p_el = para._element

    hyperlink = etree.SubElement(p_el, qn("w:hyperlink"))
    hyperlink.set(qn("w:anchor"), ref_bookmark_name(cite_key))

    run_el = etree.SubElement(hyperlink, qn("w:r"))
    rpr = etree.SubElement(run_el, qn("w:rPr"))
    rstyle = etree.SubElement(rpr, qn("w:rStyle"))
    rstyle.set(qn("w:val"), CITATION_REF_STYLE)

    t_el = etree.SubElement(run_el, qn("w:t"))
    t_el.text = text
    t_el.set(qn("xml:space"), "preserve")


def _next_bookmark_id(doc) -> int:
    max_id = 0
    for bm in doc.element.body.iter(qn("w:bookmarkStart")):
        try:
            bm_id = int(bm.get(qn("w:id"), "0"))
            if bm_id > max_id:
                max_id = bm_id
        except ValueError:
            pass
    return max_id + 1


def _add_ref_bookmark(para, key: str, doc) -> None:
    p_el = para._element
    bm_id = str(_next_bookmark_id(doc))
    name = ref_bookmark_name(key)

    bm_start = etree.Element(qn("w:bookmarkStart"))
    bm_start.set(qn("w:id"), bm_id)
    bm_start.set(qn("w:name"), name)

    ppr = p_el.find(qn("w:pPr"))
    if ppr is not None:
        ppr.addnext(bm_start)
    else:
        p_el.insert(0, bm_start)

    bm_end = etree.SubElement(p_el, qn("w:bookmarkEnd"))
    bm_end.set(qn("w:id"), bm_id)
