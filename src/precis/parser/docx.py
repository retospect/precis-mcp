"""DOCX parser using python-docx."""

from __future__ import annotations

import os
import tempfile
from pathlib import Path

from docx import Document
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.ns import qn
from lxml import etree

from precis.citations import (
    BIB_DEF_RE,
    BIB_ENTRY_STYLE,
    CITATION_REF_STYLE,
    ORPHAN_PREFIX,
    CitationIndex,
    parse_ref_bookmark,
    ref_bookmark_name,
)
from precis.formatting import FormattedRun, markdown_to_runs, runs_to_markdown
from precis.nodes import Node, PathCounter, make_slug, resolve_slug


class DocxParser:
    """Parse and manipulate .docx documents."""

    def parse(self, path: Path) -> list[Node]:
        """Parse a DOCX file into a list of Nodes."""
        doc = Document(str(path))
        comments_map = _parse_comments(doc)
        counter = PathCounter()
        nodes: list[Node] = []
        slug_counts: dict[str, int] = {}

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                para = None
                for p in doc.paragraphs:
                    if p._element is element:
                        para = p
                        break
                if para is None:
                    continue

                text = para.text.strip()
                if not text:
                    continue

                style_name = para.style.name if para.style else "Normal"

                # Check if heading
                heading_level = _heading_level(style_name)
                if heading_level:
                    node_path = counter.next_heading(heading_level)
                    base_slug = make_slug(text)
                    slug = resolve_slug(base_slug, slug_counts)
                    nodes.append(
                        Node(
                            slug=slug,
                            path=node_path,
                            node_type="h",
                            text=text,
                            precis=text,
                            style=style_name,
                            comments=_collect_para_comments(element, comments_map),
                        )
                    )
                elif style_name == BIB_ENTRY_STYLE:
                    # Bibliography entry
                    node_path = counter.next_child("b")
                    md_text = _para_to_markdown(para)
                    label = _extract_bib_label(element)
                    base_slug = make_slug(md_text)
                    slug = resolve_slug(base_slug, slug_counts)
                    nodes.append(
                        Node(
                            slug=slug,
                            path=node_path,
                            node_type="b",
                            text=md_text,
                            precis=_bib_precis(label, text),
                            style=style_name,
                            label=label,
                            comments=_collect_para_comments(element, comments_map),
                        )
                    )
                else:
                    node_path = counter.next_child("p")
                    md_text = _para_to_markdown(para)
                    base_slug = make_slug(md_text)
                    slug = resolve_slug(base_slug, slug_counts)
                    nodes.append(
                        Node(
                            slug=slug,
                            path=node_path,
                            node_type="p",
                            text=md_text,
                            style=style_name,
                            comments=_collect_para_comments(element, comments_map),
                        )
                    )

            elif tag == "tbl":
                table = None
                for t in doc.tables:
                    if t._element is element:
                        table = t
                        break
                if table is None:
                    continue

                md_table = _table_to_markdown(table)
                synopsis = _table_synopsis(table)
                node_path = counter.next_child("t")
                base_slug = make_slug(md_table)
                slug = resolve_slug(base_slug, slug_counts)
                nodes.append(
                    Node(
                        slug=slug,
                        path=node_path,
                        node_type="t",
                        text=md_table,
                        precis=synopsis,
                        style="Table",
                    )
                )

        return nodes

    def source_files(self, path: Path) -> list[Path]:
        return [path]

    def write_node(self, path: Path, node: Node, new_text: str) -> None:
        """Replace a node's text content."""
        doc = Document(str(path))
        element = self._find_element(doc, node)
        if element is None:
            raise ValueError(f"Node not found in document: {node.slug}")

        if node.node_type == "p" or node.node_type == "h":
            para = self._element_to_para(doc, element)
            if para is None:
                raise ValueError(f"Paragraph not found: {node.slug}")
            _set_para_text(para, new_text, node.node_type == "h")

        self._atomic_save(doc, path)

    def insert_after(
        self, path: Path, anchor: Node, new_text: str, heading_level: int = 0
    ) -> None:
        doc = Document(str(path))
        element = self._find_element(doc, anchor)
        if element is None:
            raise ValueError(f"Anchor not found: {anchor.slug}")

        new_para = _make_paragraph(doc, new_text, heading_level)
        element.addnext(new_para._element)
        self._atomic_save(doc, path)

    def insert_before(
        self, path: Path, anchor: Node, new_text: str, heading_level: int = 0
    ) -> None:
        doc = Document(str(path))
        element = self._find_element(doc, anchor)
        if element is None:
            raise ValueError(f"Anchor not found: {anchor.slug}")

        new_para = _make_paragraph(doc, new_text, heading_level)
        element.addprevious(new_para._element)
        self._atomic_save(doc, path)

    def delete_node(self, path: Path, node: Node) -> None:
        doc = Document(str(path))
        element = self._find_element(doc, node)
        if element is None:
            raise ValueError(f"Node not found: {node.slug}")
        element.getparent().remove(element)
        self._atomic_save(doc, path)

    def append_node(self, path: Path, new_text: str, heading_level: int = 0) -> None:
        doc = Document(str(path))
        new_para = _make_paragraph(doc, new_text, heading_level)
        doc.element.body.append(new_para._element)
        self._atomic_save(doc, path)

    def move_nodes(self, path: Path, nodes: list[Node], after: Node) -> None:
        doc = Document(str(path))
        after_el = self._find_element(doc, after)
        if after_el is None:
            raise ValueError(f"Target not found: {after.slug}")

        # Collect elements to move
        elements = []
        for n in nodes:
            el = self._find_element(doc, n)
            if el is not None:
                elements.append(el)

        # Remove from current positions
        for el in elements:
            el.getparent().remove(el)

        # Insert after target in order
        insert_point = after_el
        for el in elements:
            insert_point.addnext(el)
            insert_point = el

        self._atomic_save(doc, path)

    def write_tracked(
        self, path: Path, node: Node, new_text: str, author: str = "precis"
    ) -> None:
        """Replace node text with track-changes markup."""
        doc = Document(str(path))
        element = self._find_element(doc, node)
        if element is None:
            raise ValueError(f"Node not found: {node.slug}")

        if node.node_type in ("p", "h"):
            para = self._element_to_para(doc, element)
            if para is None:
                raise ValueError(f"Paragraph not found: {node.slug}")
            _inject_tracked_replace(para, new_text, author)

        self._atomic_save(doc, path)

    def write_comment(
        self, path: Path, node: Node, text: str, author: str = "precis"
    ) -> int:
        """Add a Word comment (margin annotation) on a node. Returns comment ID."""
        doc = Document(str(path))
        element = self._find_element(doc, node)
        if element is None:
            raise ValueError(f"Node not found: {node.slug}")

        comment_id = _inject_comment(doc, element, text, author)
        self._atomic_save(doc, path)
        return comment_id

    def _find_element(self, doc: Document, node: Node):
        """Find the XML element matching a node by reparsing and matching."""
        counter = PathCounter()
        slug_counts: dict[str, int] = {}

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                para = self._element_to_para(doc, element)
                if para is None or not para.text.strip():
                    continue
                style_name = para.style.name if para.style else "Normal"
                heading_level = _heading_level(style_name)
                if heading_level:
                    counter.next_heading(heading_level)
                    text = para.text.strip()
                elif style_name == BIB_ENTRY_STYLE:
                    counter.next_child("b")
                    text = _para_to_markdown(para)
                else:
                    counter.next_child("p")
                    text = _para_to_markdown(para)
                base_slug = make_slug(text)
                slug = resolve_slug(base_slug, slug_counts)
                if slug == node.slug:
                    return element

            elif tag == "tbl":
                table = self._element_to_table(doc, element)
                if table is None:
                    continue
                counter.next_child("t")
                md = _table_to_markdown(table)
                base_slug = make_slug(md)
                slug = resolve_slug(base_slug, slug_counts)
                if slug == node.slug:
                    return element

        return None

    def _element_to_para(self, doc, element):
        for p in doc.paragraphs:
            if p._element is element:
                return p
        return None

    def _element_to_table(self, doc, element):
        for t in doc.tables:
            if t._element is element:
                return t
        return None

    def _atomic_save(self, doc: Document, path: Path) -> None:
        fd, tmp = tempfile.mkstemp(dir=path.parent, suffix=".docx.tmp")
        os.close(fd)
        try:
            doc.save(tmp)
            os.replace(tmp, path)
        except Exception:
            try:
                os.unlink(tmp)
            except OSError:
                pass
            raise


def _heading_level(style_name: str) -> int:
    """Extract heading level from style name, or 0."""
    if not style_name:
        return 0
    s = style_name.lower()
    if s.startswith("heading"):
        rest = s[7:].strip()
        if rest.isdigit():
            level = int(rest)
            if 1 <= level <= 4:
                return level
    return 0


def _para_to_markdown(para) -> str:
    """Convert a python-docx paragraph to Markdown text, including citations."""
    runs_list: list[FormattedRun] = []

    for child in para._element:
        tag = _local_tag(child)

        if tag == "hyperlink":
            anchor = child.get(qn("w:anchor"), "")
            cite_key = ""
            if anchor.startswith("ref_"):
                cite_key = anchor[4:]  # strip "ref_" prefix

            for r_el in child.findall(qn("w:r")):
                text = _run_element_text(r_el)
                if not text:
                    continue
                if cite_key:
                    runs_list.append(FormattedRun(text=text, cite_key=cite_key))
                else:
                    runs_list.append(_run_element_to_formatted(r_el, text))

        elif tag == "r":
            text = _run_element_text(child)
            if not text:
                continue
            # Check for CitationRef style without hyperlink (orphaned paste)
            if _has_char_style(child, CITATION_REF_STYLE):
                orphan_key = f"{ORPHAN_PREFIX}{text}"
                runs_list.append(FormattedRun(text=text, cite_key=orphan_key))
            else:
                runs_list.append(_run_element_to_formatted(child, text))

    return runs_to_markdown(runs_list)


def _set_para_text(
    para, markdown_text: str, is_heading: bool = False, doc=None
) -> None:
    """Set paragraph text from Markdown, creating appropriate runs."""
    # Clear existing runs and hyperlinks
    for child in list(para._element):
        tag = _local_tag(child)
        if tag in ("r", "hyperlink"):
            para._element.remove(child)

    if is_heading:
        # Strip # prefix if present
        text = markdown_text.lstrip("#").strip()
        run = para.add_run(text)
        return

    formatted_runs = markdown_to_runs(markdown_text)
    _write_formatted_runs(para, formatted_runs, doc)


def _write_formatted_runs(para, formatted_runs: list[FormattedRun], doc=None) -> None:
    """Write FormattedRun objects to a paragraph, handling citations as hyperlinks."""
    for fr in formatted_runs:
        if fr.cite_key and doc is not None:
            _add_citation_hyperlink(para, fr.text, fr.cite_key, doc)
        else:
            run = para.add_run(fr.text)
            if fr.bold:
                run.bold = True
            if fr.italic:
                run.italic = True
            if fr.superscript:
                run.font.superscript = True
            if fr.subscript:
                run.font.subscript = True
            if fr.strike:
                run.font.strike = True


def _make_paragraph(doc, text: str, heading_level: int = 0):
    """Create a new paragraph with optional heading style."""
    # Check for bibliography definition: [@key]: reference text
    bib_m = BIB_DEF_RE.match(text)
    if bib_m:
        key = bib_m.group(1)
        ref_text = bib_m.group(2)
        _ensure_styles(doc)
        para = doc.add_paragraph()
        para.style = doc.styles[BIB_ENTRY_STYLE]
        # Write the reference text as normal runs
        formatted_runs = markdown_to_runs(ref_text)
        _write_formatted_runs(para, formatted_runs, doc)
        # Add ref_ bookmark around the content
        _add_ref_bookmark(para, key, doc)
        return para

    # Check if text starts with # for heading
    if text.startswith("#"):
        level = 0
        stripped = text
        while stripped.startswith("#"):
            level += 1
            stripped = stripped[1:]
        stripped = stripped.strip()
        if level >= 1:
            heading_level = min(level, 4)
            text = stripped

    if heading_level:
        para = doc.add_paragraph(text)
        para.style = doc.styles[f"Heading {heading_level}"]
    else:
        para = doc.add_paragraph()
        formatted_runs = markdown_to_runs(text)
        _write_formatted_runs(para, formatted_runs, doc)

    return para


def _table_to_markdown(table) -> str:
    """Convert a python-docx table to Markdown pipe table."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")

    if len(rows) >= 1:
        # Add separator after header
        ncols = len(table.rows[0].cells)
        sep = "| " + " | ".join(["---"] * ncols) + " |"
        rows.insert(1, sep)

    return "\n".join(rows)


def _table_synopsis(table) -> str:
    """Generate table precis: column headers + row count + first row sample."""
    if not table.rows:
        return "[empty table]"
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    nrows = len(table.rows) - 1  # exclude header
    synopsis = "|".join(headers) + f" ({nrows}r)"
    return synopsis


def _inject_tracked_replace(para, new_text: str, author: str) -> None:
    """Inject Word track-changes markup for a paragraph replacement."""
    from datetime import datetime, timezone

    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    nsmap = para._element.nsmap
    w = nsmap.get("w", "http://schemas.openxmlformats.org/wordprocessingml/2006/main")

    # Wrap existing runs in <w:del>
    for run_el in list(para._element.findall(qn("w:r"))):
        del_el = etree.SubElement(para._element, qn("w:del"))
        del_el.set(qn("w:author"), author)
        del_el.set(qn("w:date"), now)
        para._element.remove(run_el)
        # Change w:t to w:delText
        for t_el in run_el.findall(qn("w:t")):
            t_el.tag = qn("w:delText")
            t_el.set(qn("xml:space"), "preserve")
        del_el.append(run_el)

    # Add new text as <w:ins>
    ins_el = etree.SubElement(para._element, qn("w:ins"))
    ins_el.set(qn("w:author"), author)
    ins_el.set(qn("w:date"), now)

    formatted_runs = markdown_to_runs(new_text)
    for fr in formatted_runs:
        run_el = etree.SubElement(ins_el, qn("w:r"))
        t_el = etree.SubElement(run_el, qn("w:t"))
        t_el.text = fr.text
        t_el.set(qn("xml:space"), "preserve")

        if fr.bold or fr.italic or fr.superscript or fr.subscript or fr.strike:
            rpr = etree.SubElement(run_el, qn("w:rPr"))
            run_el.insert(0, rpr)
            if fr.bold:
                etree.SubElement(rpr, qn("w:b"))
            if fr.italic:
                etree.SubElement(rpr, qn("w:i"))
            if fr.strike:
                etree.SubElement(rpr, qn("w:strike"))
            if fr.superscript:
                vertAlign = etree.SubElement(rpr, qn("w:vertAlign"))
                vertAlign.set(qn("w:val"), "superscript")
            if fr.subscript:
                vertAlign = etree.SubElement(rpr, qn("w:vertAlign"))
                vertAlign.set(qn("w:val"), "subscript")


# ---------------------------------------------------------------------------
# Word comments (margin annotations)
# ---------------------------------------------------------------------------

_COMMENTS_URI = PackURI("/word/comments.xml")
_RT_COMMENTS = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
_COMMENTS_CT = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)


def _get_comments_part(doc):
    """Get existing comments part, or None."""
    for rel in doc.part.rels.values():
        if rel.reltype == _RT_COMMENTS:
            return rel.target_part
    return None


def _get_or_create_comments_part(doc):
    """Get or create the comments.xml part."""
    part = _get_comments_part(doc)
    if part is not None:
        return part
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    root = etree.Element(qn("w:comments"), nsmap=nsmap)
    blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    part = Part(_COMMENTS_URI, _COMMENTS_CT, blob, doc.part.package)
    doc.part.relate_to(part, _RT_COMMENTS)
    return part


def _parse_comments(doc) -> dict[int, dict]:
    """Parse comments part → {id: {author, text}}."""
    part = _get_comments_part(doc)
    if part is None:
        return {}
    root = etree.fromstring(part.blob)
    result = {}
    for comment_el in root.findall(qn("w:comment")):
        try:
            cid = int(comment_el.get(qn("w:id"), "0"))
        except ValueError:
            continue
        author = comment_el.get(qn("w:author"), "")
        texts = []
        for p_el in comment_el.findall(qn("w:p")):
            for r_el in p_el.findall(qn("w:r")):
                for t_el in r_el.findall(qn("w:t")):
                    if t_el.text:
                        texts.append(t_el.text)
        result[cid] = {"author": author, "text": " ".join(texts)}
    return result


def _next_comment_id(doc) -> int:
    """Find next available comment ID."""
    part = _get_comments_part(doc)
    if part is None:
        return 1
    root = etree.fromstring(part.blob)
    max_id = 0
    for comment_el in root.findall(qn("w:comment")):
        try:
            cid = int(comment_el.get(qn("w:id"), "0"))
            if cid > max_id:
                max_id = cid
        except ValueError:
            pass
    return max_id + 1


def _collect_para_comments(element, comments_map: dict) -> list[dict]:
    """Collect comments attached to a paragraph element."""
    comments = []
    seen_ids: set[int] = set()
    # commentRangeStart as direct child
    for el in element.findall(qn("w:commentRangeStart")):
        try:
            cid = int(el.get(qn("w:id"), "0"))
        except ValueError:
            continue
        if cid in comments_map and cid not in seen_ids:
            seen_ids.add(cid)
            cm = comments_map[cid]
            comments.append({"id": cid, "author": cm["author"], "text": cm["text"]})
    # commentReference in runs (fallback for Word-created comments)
    for r_el in element.findall(qn("w:r")):
        for ref_el in r_el.findall(qn("w:commentReference")):
            try:
                cid = int(ref_el.get(qn("w:id"), "0"))
            except ValueError:
                continue
            if cid in comments_map and cid not in seen_ids:
                seen_ids.add(cid)
                cm = comments_map[cid]
                comments.append({"id": cid, "author": cm["author"], "text": cm["text"]})
    return comments


def _inject_comment(doc, para_element, text: str, author: str = "precis") -> int:
    """Add a Word comment on a paragraph. Returns the comment ID."""
    from datetime import datetime, timezone

    comment_id = _next_comment_id(doc)
    now = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    cid_str = str(comment_id)

    # Add comment to comments part
    comments_part = _get_or_create_comments_part(doc)

    # XmlPart (loaded from file) stores _element; regular Part uses _blob.
    # XmlPart.blob serializes from _element, ignoring _blob — so we must
    # modify _element in-place when it exists.
    if hasattr(comments_part, "_element") and comments_part._element is not None:
        root = comments_part._element
    else:
        root = etree.fromstring(comments_part.blob)

    comment_el = etree.SubElement(root, qn("w:comment"))
    comment_el.set(qn("w:id"), cid_str)
    comment_el.set(qn("w:author"), author)
    comment_el.set(qn("w:date"), now)
    comment_el.set(qn("w:initials"), author[:2].upper())

    p_el = etree.SubElement(comment_el, qn("w:p"))
    r_el = etree.SubElement(p_el, qn("w:r"))
    t_el = etree.SubElement(r_el, qn("w:t"))
    t_el.text = text
    t_el.set(qn("xml:space"), "preserve")

    # Write back only for regular Part (XmlPart auto-serializes from _element)
    if not hasattr(comments_part, "_element") or comments_part._element is None:
        comments_part._blob = etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

    # Add range markers to paragraph
    range_start = etree.Element(qn("w:commentRangeStart"))
    range_start.set(qn("w:id"), cid_str)

    ppr = para_element.find(qn("w:pPr"))
    first_run = para_element.find(qn("w:r"))
    if first_run is not None:
        first_run.addprevious(range_start)
    elif ppr is not None:
        ppr.addnext(range_start)
    else:
        para_element.insert(0, range_start)

    # commentRangeEnd after last content
    range_end = etree.SubElement(para_element, qn("w:commentRangeEnd"))
    range_end.set(qn("w:id"), cid_str)

    # Comment reference run
    ref_run = etree.SubElement(para_element, qn("w:r"))
    ref_el = etree.SubElement(ref_run, qn("w:commentReference"))
    ref_el.set(qn("w:id"), cid_str)

    return comment_id


# ---------------------------------------------------------------------------
# XML helpers for citation reading
# ---------------------------------------------------------------------------


def _local_tag(el) -> str:
    """Get the local tag name (strip namespace)."""
    tag = el.tag
    if "}" in tag:
        return tag.split("}")[-1]
    return tag


def _run_element_text(r_el) -> str:
    """Extract text from a w:r element."""
    parts = []
    for t_el in r_el.findall(qn("w:t")):
        if t_el.text:
            parts.append(t_el.text)
    return "".join(parts)


def _has_char_style(r_el, style_name: str) -> bool:
    """Check if a run element has a specific character style."""
    rpr = r_el.find(qn("w:rPr"))
    if rpr is None:
        return False
    rstyle = rpr.find(qn("w:rStyle"))
    if rstyle is None:
        return False
    return rstyle.get(qn("w:val")) == style_name


def _run_element_to_formatted(r_el, text: str) -> FormattedRun:
    """Convert a w:r XML element to a FormattedRun."""
    rpr = r_el.find(qn("w:rPr"))
    bold = False
    italic = False
    superscript = False
    subscript = False
    strike = False

    if rpr is not None:
        bold = rpr.find(qn("w:b")) is not None
        italic = rpr.find(qn("w:i")) is not None
        strike = rpr.find(qn("w:strike")) is not None
        vert = rpr.find(qn("w:vertAlign"))
        if vert is not None:
            val = vert.get(qn("w:val"), "")
            superscript = val == "superscript"
            subscript = val == "subscript"

    return FormattedRun(
        text=text,
        bold=bold,
        italic=italic,
        superscript=superscript,
        subscript=subscript,
        strike=strike,
    )


def _extract_bib_label(p_el) -> str:
    """Extract citation key from ref_ bookmark on this paragraph."""
    for bm in p_el.findall(qn("w:bookmarkStart")):
        name = bm.get(qn("w:name"), "")
        key = parse_ref_bookmark(name)
        if key:
            return key
    return ""


def _bib_precis(label: str, text: str) -> str:
    """Generate short precis for a bibliography entry."""
    short = text[:60].rstrip()
    if len(text) > 60:
        short += "…"
    if label:
        return f"{label}: {short}"
    return short


# ---------------------------------------------------------------------------
# DOCX style and citation writing helpers
# ---------------------------------------------------------------------------


def _ensure_styles(doc) -> None:
    """Create CitationRef and BibEntry styles if they don't exist."""
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Inches

    styles = doc.styles

    try:
        styles[CITATION_REF_STYLE]
    except KeyError:
        styles.add_style(CITATION_REF_STYLE, WD_STYLE_TYPE.CHARACTER)

    try:
        styles[BIB_ENTRY_STYLE]
    except KeyError:
        style = styles.add_style(BIB_ENTRY_STYLE, WD_STYLE_TYPE.PARAGRAPH)
        style.paragraph_format.first_line_indent = Inches(-0.5)
        style.paragraph_format.left_indent = Inches(0.5)


def _add_citation_hyperlink(para, text: str, cite_key: str, doc) -> None:
    """Add a citation as an internal hyperlink to ref_{key} with CitationRef style."""
    _ensure_styles(doc)
    p_el = para._element

    hyperlink = etree.SubElement(p_el, qn("w:hyperlink"))
    hyperlink.set(qn("w:anchor"), ref_bookmark_name(cite_key))

    run_el = etree.SubElement(hyperlink, qn("w:r"))
    rpr = etree.SubElement(run_el, qn("w:rPr"))
    rstyle = etree.SubElement(rpr, qn("w:rStyle"))
    rstyle.set(qn("w:val"), CITATION_REF_STYLE)

    t_el = etree.SubElement(run_el, qn("w:t"))
    t_el.text = text
    t_el.set(qn("xml:space"), "preserve")


def _next_bookmark_id(doc) -> int:
    """Find the next available bookmark ID in the document."""
    max_id = 0
    for bm in doc.element.body.iter(qn("w:bookmarkStart")):
        try:
            bm_id = int(bm.get(qn("w:id"), "0"))
            if bm_id > max_id:
                max_id = bm_id
        except ValueError:
            pass
    return max_id + 1


def _add_ref_bookmark(para, key: str, doc) -> None:
    """Wrap the paragraph content in a ref_{key} bookmark."""
    p_el = para._element
    bm_id = str(_next_bookmark_id(doc))
    name = ref_bookmark_name(key)

    # Insert bookmarkStart as first child (after pPr if present)
    bm_start = etree.Element(qn("w:bookmarkStart"))
    bm_start.set(qn("w:id"), bm_id)
    bm_start.set(qn("w:name"), name)

    ppr = p_el.find(qn("w:pPr"))
    if ppr is not None:
        ppr.addnext(bm_start)
    else:
        p_el.insert(0, bm_start)

    # Append bookmarkEnd as last child
    bm_end = etree.SubElement(p_el, qn("w:bookmarkEnd"))
    bm_end.set(qn("w:id"), bm_id)
