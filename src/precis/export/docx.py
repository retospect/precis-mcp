"""Draft → ``.docx`` export (python-docx). Sibling of ``export/latex.py``.

Synchronous and **toolchain-free** (python-docx + lxml, both already
deps) — this is the "just works" path: no latexmk, nothing to install,
nothing to compile. It mirrors the LaTeX exporter's chunk walk and inline
grammar so the two never drift, and — crucially — resolves citations
through the **same** paper-ref lookup the ``.bib`` path uses
(:func:`precis.export.latex.build_bib`). A docx and a PDF therefore cite
the *identical* resolved references: that shared resolver is the
citation-integrity guarantee.

Citation model: each ``[§slug~n]`` / ``paper:slug~n`` (and a verbatim
LaTeX ``\\cite{slug}``) becomes a superscript numbered marker ``[n]`` in
the text, backed by a numbered **References** section at the document end
carrying the resolved bibliographic entry (authors · year · title ·
DOI/arXiv). The marker is a plain run, so a paper cited many times reuses
its number at every site — deliberately *not* a native Word endnote field,
which must be referenced exactly once (reusing one makes Word declare the
file's content unreadable). Math is native OMML (see ``_render_math``).
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from precis.export.latex import _COMBINED, _bibtex_authors, preprocess_draft_inline
from precis.utils import handle_registry
from precis.utils.authors import build_byline
from precis.utils.draft_markup import DRAFT_CITE_PATTERN

#: chunk depth → Word heading level (1..4); deeper collapses to 4.
_MAX_HEADING_LEVEL = 4

#: A sober scientific-manuscript theme layered over python-docx's (bright,
#: blue-accented Calibri) default template. Times New Roman, all black, at
#: the sizes a journal submission expects. ``_apply_paper_theme`` writes
#: these onto the built-in named styles, so every paragraph/heading picks
#: them up without any per-run work. Tweak the constants here to re-tune.
_BODY_FONT = "Times New Roman"  # the journal / manuscript standard serif
_HEADING_FONT = "Times New Roman"
_BODY_PT = 12  # 12 pt body — standard for a double-spaced manuscript
_INK = "000000"  # everything black — no grey, no accent colour
#: hyperlink colour — black too (kept underlined so it still reads as a link).
_LINK_INK = "000000"
#: page margins, inches — the standard 1 in on all four sides.
_MARGIN_IN = 1.0


def _apply_paper_theme(doc: Any) -> None:
    """Reface the built-in styles + page margins to a sober paper look.

    python-docx's default template ships bright Calibri body text, blue
    Calibri-Light headings, and Word's stock margins; we overwrite the
    named styles the walk actually uses (``Normal``, ``Title``, ``Heading
    1``–``4``) with black Times New Roman and set 1-inch margins. Guarded
    per style — a template missing one of these names just skips it."""
    from docx.shared import Inches, Pt, RGBColor

    def _restyle(name: str, *, font: str, size_pt: int | None = None) -> None:
        try:
            style = doc.styles[name]
        except KeyError:  # pragma: no cover - depends on the docx template
            return
        style.font.name = font
        style.font.color.rgb = RGBColor.from_string(_INK)
        if size_pt is not None:
            style.font.size = Pt(size_pt)

    _restyle("Normal", font=_BODY_FONT, size_pt=_BODY_PT)
    _restyle("Title", font=_HEADING_FONT)
    for lvl in range(1, _MAX_HEADING_LEVEL + 1):
        _restyle(f"Heading {lvl}", font=_HEADING_FONT)

    margin = Inches(_MARGIN_IN)
    for section in doc.sections:
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


@dataclass
class DocxResult:
    """Path written + diagnostics (parallel to the LaTeX ``ExportResult``)."""

    path: Path
    cited_slugs: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


@dataclass
class _Ctx:
    store: Any
    known_handles: set[str]
    abbrevs: dict[str, str] = field(default_factory=dict)  # short → long
    cited: list[str] = field(default_factory=list)  # paper slug order = ref number
    warnings: list[str] = field(default_factory=list)
    seen_acr: set[str] = field(default_factory=set)  # already expanded once
    used_acr: set[str] = field(default_factory=set)  # for the acronyms list
    last_cite: str | None = None  # paper of the immediately-preceding mark
    endnote: bool = False  # emit EndNote CWYW fields instead of plain [n]
    resolved: dict[str, dict[str, Any]] = field(default_factory=dict)  # slug→record

    def cite_number(self, slug: str) -> int:
        """1-based reference number for a slug (stable insertion order)."""
        if slug not in self.cited:
            self.cited.append(slug)
        return self.cited.index(slug) + 1

    _short_re: Any = None

    def short_pattern(self) -> Any:
        """A compiled regex matching any known short as a whole word
        (optional trailing ``s`` for plurals), longest-first so ``MOFs``
        wins over ``MOF``. ``None`` when the draft defines no abbreviations."""
        if self._short_re is None and self.abbrevs:
            shorts = sorted(self.abbrevs, key=len, reverse=True)
            self._short_re = re.compile(
                r"\b(" + "|".join(re.escape(s) for s in shorts) + r")(s?)\b"
            )
        return self._short_re


def _add_hyperlink(paragraph: Any, url: str, text: str) -> None:
    """Append a real external hyperlink run to ``paragraph`` (python-docx
    has no native helper). Styled blue + underlined so it reads as a link.
    Used for a ROR affiliation id."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _LINK_INK)
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def _render_byline(doc: Any, byline: dict[str, Any]) -> None:
    """Emit the author byline + affiliation block under the title heading.

    Names on one line (superscript marks when >1 affiliation), then one
    affiliation per line (``¹ Org``, org hyperlinked to its ROR id). A no-
    author draft renders nothing, matching the legacy title-only output."""
    authors = byline.get("authors") or []
    if not authors:
        return
    multi = byline.get("multi")
    names_p = doc.add_paragraph()
    for i, a in enumerate(authors):
        if i:
            names_p.add_run(", ")
        names_p.add_run(a["name"])
        if multi and a.get("sup"):
            sup = names_p.add_run(a["sup"])
            sup.font.superscript = True
    for aff in byline.get("affiliations") or []:
        p = doc.add_paragraph()
        if multi:
            mark = p.add_run(str(aff["index"]) + " ")
            mark.font.superscript = True
        org = aff.get("org") or aff.get("ror") or ""
        if aff.get("ror"):
            _add_hyperlink(p, aff["ror"], org)
        else:
            run = p.add_run(org)
            run.italic = True


def export_docx(
    store: Any, ref: Any, *, target_path: Path, citations: str = "plain"
) -> DocxResult:
    """Render a draft into ``target_path`` as a ``.docx``. Returns the
    path plus the cited slugs and any resolution warnings.

    ``citations`` selects how in-text references render:

    * ``"plain"`` (default) — a superscript ``[n]`` marker + a plain
      numbered **References** section. Universal, no add-in needed.
    * ``"endnote"`` — native EndNote *Cite While You Write* fields
      (``ADDIN EN.CITE`` + ``EN.REFLIST``), so EndNote recognizes and can
      reformat / manage the citations (see :mod:`precis.export.endnote`).
    """
    from docx import Document

    target_path = Path(target_path)
    chunks = store.reading_order(ref.id)
    handles = {c.handle for c in chunks}
    ctx = _Ctx(
        store=store,
        known_handles=handles,
        abbrevs=store.defined_abbrevs(ref.id),
        endnote=(citations == "endnote"),
    )

    doc = Document()
    _apply_paper_theme(doc)
    terms = store.draft_terms(ref.id)  # handle → (short, long)
    byline = build_byline(getattr(ref, "authors", None))

    # List context (migration 0037): a ulist/olist container owns `item`
    # children, which render with Word's built-in List Bullet / List Number
    # styles. The immediate container decides bullet vs number; the count of
    # list-container ancestors picks the nesting level (1–3).
    _kind_by_id = {c.chunk_id: c.chunk_kind for c in chunks}
    _parent_by_id = {c.chunk_id: c.parent_chunk_id for c in chunks}

    # The first heading at depth 0 is the title — render it as the doc title.
    title_done = False
    for c in chunks:
        kind = c.chunk_kind
        if kind in ("ulist", "olist"):
            continue  # structural container — its items carry the prose
        if kind == "item":
            base, level = "List Bullet", 0
            pid = c.parent_chunk_id
            while pid is not None:
                pk = _kind_by_id.get(pid)
                if pk in ("ulist", "olist"):
                    if level == 0:
                        base = "List Number" if pk == "olist" else "List Bullet"
                    level += 1
                pid = _parent_by_id.get(pid)
            style = base if level <= 1 else f"{base} {min(level, 3)}"
            p = doc.add_paragraph(style=style)
            _render_inline(c.text, ctx, p)
            continue
        if kind == "term":
            # Render in place (terms live under the draft's own Glossary
            # heading) as "SHORT — long", pulling the short from meta.
            short, long = terms.get(c.handle, ("", c.text))
            p = doc.add_paragraph()
            if short:
                p.add_run(short).bold = True
                p.add_run(f" — {long}")
            else:
                p.add_run(long)
            continue
        if kind == "heading":
            if not title_done and c.depth == 0:
                doc.add_heading(c.text, level=0)
                _render_byline(doc, byline)
                title_done = True
                continue
            level = min(max(c.depth, 1), _MAX_HEADING_LEVEL)
            doc.add_heading(c.text, level=level)
            continue
        if kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(c.text)
            run.font.name = "Consolas"
            continue
        if kind == "table":
            _render_table(doc, c, ctx)
            continue
        # paragraph (default)
        p = doc.add_paragraph()
        _render_inline(c.text, ctx, p)

    # An authored glossary (``term`` chunks, rendered in place above) already
    # lists the abbreviations, so the auto "Acronyms" section would duplicate
    # it — only emit it when the draft defines no terms of its own.
    if not terms:
        _append_acronyms(doc, ctx)
    _append_references(doc, ctx)
    if ctx.endnote and ctx.cited:
        from precis.export.endnote import install_document_vars

        install_document_vars(
            doc, list(range(1, len(ctx.cited) + 1)), style="Annotated"
        )

    target_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target_path))
    return DocxResult(
        path=target_path, cited_slugs=list(ctx.cited), warnings=ctx.warnings
    )


# ── table rendering ───────────────────────────────────────────────


def _render_table(doc: Any, chunk: Any, ctx: _Ctx) -> None:
    """Render a ``chunk_kind='table'`` chunk as a native Word table (ADR
    0035 §1). The canonical ``meta.table`` is recovered via the shared
    :func:`precis.utils.table_data.table_payload`; cells go through the same
    inline grammar as prose so citations / math / abbreviations inside a
    cell resolve identically. Falls back to a plain paragraph when the table
    can't be recovered."""
    from precis.utils.table_data import table_payload

    payload = table_payload(getattr(chunk, "meta", None), chunk.text)
    if payload is None:
        _render_inline(chunk.text, ctx, doc.add_paragraph())
        return
    header, rows, caption = payload["header"], payload["rows"], payload["caption"]
    if caption:
        cap = doc.add_paragraph()
        _render_inline(caption, ctx, cap)
        for run in cap.runs:
            run.bold = True
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    try:  # built-in style; absent in some templates → fall back to no style
        table.style = "Table Grid"
    except KeyError:  # pragma: no cover - depends on the docx template
        pass
    for cell, text in zip(table.rows[0].cells, header):
        para = cell.paragraphs[0]
        para.add_run(text).bold = True
    for row in rows:
        cells = table.add_row().cells
        for j in range(ncols):
            _render_inline(row[j] if j < len(row) else "", ctx, cells[j].paragraphs[0])


# ── inline rendering ──────────────────────────────────────────────


def _render_inline(text: str, ctx: _Ctx, paragraph: Any) -> None:
    """Walk a chunk's text, adding runs to ``paragraph``. References go
    through :func:`_render_reference`; the prose gaps between them get
    markdown/sub-sup/math run formatting."""
    text = preprocess_draft_inline(text)
    last = 0
    for m in _COMBINED.finditer(text):
        _render_gap(text[last : m.start()], ctx, paragraph)
        _render_reference(m, ctx, paragraph)
        last = m.end()
    _render_gap(text[last:], ctx, paragraph)


# Tokeniser for a non-reference gap: math / code / bold / italic / sub /
# sup become typed spans; everything else is plain text. Ordered so the
# verbatim spans (math, code) are carved out before emphasis.
_SPAN = re.compile(
    r"(?P<math>\$\$.+?\$\$|\$[^$]+\$)"
    r"|(?P<code>`[^`]+`)"
    r"|(?P<sub><sub>.+?</sub>)"
    r"|(?P<sup><sup>.+?</sup>)"
    r"|(?P<bold>\*\*.+?\*\*)"
    r"|(?P<italic>(?<![\*\w])\*(?!\s)[^*]+?(?<!\s)\*(?!\w))",
    re.DOTALL,
)


def _render_gap(text: str, ctx: _Ctx, paragraph: Any) -> None:
    if not text:
        return
    # Real prose between two citations breaks the consecutive-cite run;
    # whitespace-only (``[§a~3] [§a~9]``) does not, so those still collapse.
    if text.strip():
        ctx.last_cite = None
    last = 0
    for m in _SPAN.finditer(text):
        if m.start() > last:
            _emit_text(text[last : m.start()], ctx, paragraph)
        if m.group("math") is not None:
            _render_math(m.group("math"), paragraph)
        elif m.group("code") is not None:
            r = paragraph.add_run(m.group("code")[1:-1])
            r.font.name = "Consolas"
        elif m.group("sub") is not None:
            r = paragraph.add_run(m.group("sub")[5:-6])
            r.font.subscript = True
        elif m.group("sup") is not None:
            r = paragraph.add_run(m.group("sup")[5:-6])
            r.font.superscript = True
        elif m.group("bold") is not None:
            paragraph.add_run(m.group("bold")[2:-2]).bold = True
        elif m.group("italic") is not None:
            paragraph.add_run(m.group("italic")[1:-1]).italic = True
        last = m.end()
    if last < len(text):
        _emit_text(text[last:], ctx, paragraph)


def _emit_text(text: str, ctx: _Ctx, paragraph: Any) -> None:
    """Plain prose → runs, with **render-time first-use expansion** of
    known abbreviations: the first occurrence (in reading order) of a
    defined short becomes ``Long Form (SHORT)``; later ones stay ``SHORT``.
    No authoring markup — this mirrors what the LaTeX ``\\gls`` path does,
    and survives chunk reordering because it's computed here at export.
    A trailing ``s`` (plural) is preserved on the short."""
    pat = ctx.short_pattern()
    if pat is None:
        paragraph.add_run(text)
        return
    last = 0
    for m in pat.finditer(text):
        if m.start() > last:
            paragraph.add_run(text[last : m.start()])
        short, plural = m.group(1), m.group(2)
        ctx.used_acr.add(short)
        if short not in ctx.seen_acr:
            ctx.seen_acr.add(short)
            paragraph.add_run(f"{ctx.abbrevs[short]} ({short}{plural})")
        else:
            paragraph.add_run(f"{short}{plural}")
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def _render_math(span: str, paragraph: Any) -> None:
    """Render ``$…$`` / ``$$…$$`` as native Word math (OMML), so equations
    typeset rather than appearing as literal source. Falls back to italic
    text when the LaTeX can't be converted (missing dep / exotic macro)."""
    inner = span.strip("$").strip()
    from precis.export.omml import latex_to_omml

    omath = latex_to_omml(inner)
    if omath is None:
        paragraph.add_run(inner).italic = True
        return
    # Append the <m:oMath> element directly into the paragraph's XML
    # (inline math sits among the runs).
    paragraph._p.append(omath)


def _render_reference(m: re.Match[str], ctx: _Ctx, paragraph: Any) -> None:
    """One matched inline reference. Citations → a numbered ``[n]``
    superscript marker (and register the slug). Cross-refs render their
    surface text; authoring links / bare thought mentions render nothing
    (provenance only) — mirrors the LaTeX exporter."""
    if m.group("auth") is not None:
        return  # [[…]] authoring link — provenance only
    if m.group("disp") is not None:
        _render_target(m.group("tgt"), m.group("disp"), ctx, paragraph)
        return
    if m.group("bare") is not None:
        _render_target(m.group("bare"), None, ctx, paragraph)
        return
    if m.group("ref") is not None:
        if m.group("kind") == "paper":
            _cite(m.group("id"), ctx, paragraph)
        return  # bare memory:/think:/… — not citeable
    if m.group("bare_conv") is not None:
        return
    if m.group("bare_paper") is not None:
        _cite(m.group("bare_paper").split("~", 1)[0], ctx, paragraph)
        return


def _handle_cite_key(tgt: str, ctx: _Ctx) -> tuple[str, int | None] | None:
    """A paper/patent handle (``pc<chunk_id>`` / ``pa<ref_id>``) → its
    ``(cite_key, chunk_id)``, via the one store resolver. ``chunk_id`` is set
    only for a chunk handle (``pc<id>``) — the specific cited passage, which
    the EndNote path embeds as a traveling note. ``None`` if it doesn't
    resolve to a live paper. Mirrors :func:`precis.export.latex._handle_cite_key`
    so the docx and PDF paths cite the identical resolved key."""
    if ctx.store is None:
        return None
    try:
        resolved = ctx.store.resolve_handle(tgt)
    except Exception:  # pragma: no cover — store hiccup
        return None
    if resolved is None or not resolved.public_id:
        return None
    return resolved.public_id, resolved.chunk_id


def _finding_cite_key(tgt: str, ctx: _Ctx) -> str | None:
    """A finding handle (``fi<id>``) → its bibliographic key (primary
    cite_key once the chase establishes it, else the ``pub_id`` stub).
    Mirrors :func:`precis.export.latex._finding_cite_key`."""
    if ctx.store is None:
        return None
    parsed = handle_registry.parse(tgt)
    if parsed is None:
        return None
    _kind, _is_chunk, pk = parsed
    ref = ctx.store.fetch_refs_by_ids([pk]).get(pk)
    if ref is None:
        return None
    meta = ref.meta or {}
    key = meta.get("primary_cite_key") or meta.get("pub_id")
    return str(key) if key else None


def _render_target(tgt: str, surface: str | None, ctx: _Ctx, paragraph: Any) -> None:
    if tgt.startswith("§"):  # a citation — keep the consecutive-cite run
        cm = DRAFT_CITE_PATTERN.fullmatch(tgt)
        if cm is not None:
            _cite(cm.group("slug"), ctx, paragraph)
        return
    # ADR 0036 universal handle: ``[pc10]`` / ``[pa5]`` (a paper) and
    # ``[pt7]`` (a patent) are citations; ``[fi3]`` a finding cite; ``[dc41]``
    # an intra-draft cross-ref; a record handle for a thought (``[me5]``) is
    # provenance-only. The LaTeX exporter resolves these — the docx path used
    # to drop them all, so handle-cited drafts (e.g. everything imported from
    # LaTeX) rendered with no citations and no References section at all.
    parsed = handle_registry.parse(tgt)
    if parsed is not None:
        kind, is_chunk, _pk = parsed
        if kind in ("paper", "patent"):
            hit = _handle_cite_key(tgt, ctx)
            if hit:
                slug, chunk_id = hit
                _cite(slug, ctx, paragraph, chunk_id=chunk_id)
            return
        if kind == "finding":
            slug = _finding_cite_key(tgt, ctx)
            if slug:
                _cite(slug, ctx, paragraph)
            return
        # draft cross-ref / other record handle → not a citation.
        ctx.last_cite = None
        if kind == "draft" and is_chunk and surface:
            paragraph.add_run(surface)  # no Word cross-ref field yet — text only
        return
    # Any non-citation content breaks a run of consecutive citations.
    ctx.last_cite = None
    if tgt.startswith("¶"):
        handle = tgt[1:]
        if handle not in ctx.known_handles:
            ctx.warnings.append(f"cross-ref ¶{handle}: no such live chunk — downgraded")
        paragraph.add_run(surface or f"¶{handle}")
        return
    if tgt.startswith(("http://", "https://")):
        paragraph.add_run(surface or tgt)
        return


def _cite(slug: str, ctx: _Ctx, paragraph: Any, chunk_id: int | None = None) -> None:
    """Emit a numbered citation marker — a superscript ``[n]`` keyed on the
    **paper**. The numbered **References** section at the document end
    (:func:`_append_references`) carries the resolved entry, so entry ``n``
    backs every in-text ``[n]``.

    Citations key on the paper, not the chunk: ``a~3``, ``a~9``, ``a~23``
    all reference paper ``a`` → mark ``[n]`` with the same ``n`` each time.
    A paper cited many times prints ``[n]`` at each site — a plain run, so
    it *repeats* freely (a real Word endnote field can't: endnotes are
    1:1 with their reference, and reusing one corrupts the document). And
    **consecutive** marks for the same paper collapse to a single mark.

    ``chunk_id`` (set for a ``pc<id>`` chunk citation) is the specific cited
    passage — carried into the EndNote field's traveling note."""
    slug = slug.split("~", 1)[0]  # the paper, not the cited chunk
    n = ctx.cite_number(slug)  # registers the paper (idempotent)
    if ctx.last_cite == slug:
        return  # consecutive cite to the same paper — one mark for the run
    ctx.last_cite = slug
    if ctx.endnote:
        _cite_endnote(slug, n, ctx, paragraph, chunk_id)
        return
    run = paragraph.add_run(f"[{n}]")
    run.font.superscript = True


def _resolve_source(store: Any, slug: str, rec_number: int) -> dict[str, Any] | None:
    """Resolve a cited slug to the structured record EndNote embeds — the
    SAME paper/patent lookup the plain path and the ``.bib`` path use, so the
    docx, PDF and EndNote outputs cite the identical resolved source."""
    pref = store.get_ref(kind="paper", id=slug) or store.get_ref(kind="patent", id=slug)
    if pref is None:
        return None
    doi = arxiv = None
    try:
        alias = store.identifiers_for_refs([pref.id]).get(pref.id, {})
        doi, arxiv = alias.get("doi"), alias.get("arxiv")
    except Exception:  # pragma: no cover — identifier lookup best-effort
        pass
    meta = pref.meta or {}
    url = (
        f"https://doi.org/{doi}"
        if doi
        else (f"https://arxiv.org/abs/{arxiv}" if arxiv else None)
    )
    return {
        "kind": pref.kind,
        "tag": slug,
        "rec_number": rec_number,
        "authors": pref.authors,
        "title": pref.title or slug,
        "year": pref.year,
        "journal": meta.get("venue")
        or meta.get("journal")
        or meta.get("container_title"),
        "volume": meta.get("volume"),
        "doi": doi,
        "url": url,
    }


def _cite_endnote(
    slug: str, n: int, ctx: _Ctx, paragraph: Any, chunk_id: int | None = None
) -> None:
    """Emit a native EndNote ``ADDIN EN.CITE`` field for one citation. An
    unresolved slug degrades to a plain superscript marker (nothing lost).

    A ``pc<id>`` chunk citation (``chunk_id`` set) embeds that chunk's text as
    the field's traveling note — so the exact cited passage rides along in
    EndNote. Per-citation, not per-paper: different cite sites of the same
    paper carry the passage each one actually pointed at."""
    from precis.export.endnote import add_citation_field

    source = ctx.resolved.get(slug)
    if source is None:
        source = _resolve_source(ctx.store, slug, n)
        if source is None:
            ctx.warnings.append(f"cite {slug!r}: no paper in corpus — plain marker")
            paragraph.add_run(f"[{n}]").font.superscript = True
            return
        ctx.resolved[slug] = source
    notes: str | None = None
    if chunk_id is not None and ctx.store is not None:
        try:
            notes = ctx.store.chunk_text_by_id(chunk_id)
        except Exception:  # pragma: no cover — chunk fetch best-effort
            notes = None
    add_citation_field(paragraph, source, notes=notes)


# ── references + glossary sections ────────────────────────────────


def _format_reference(store: Any, slug: str, warnings: list[str]) -> str:
    """One reference line, resolved through the SAME paper lookup as the
    ``.bib`` path (citation-integrity parity with the PDF). A slug with no
    paper in the corpus degrades to a marked stub + a warning."""
    pref = store.get_ref(kind="paper", id=slug) or store.get_ref(kind="patent", id=slug)
    if pref is None:
        warnings.append(f"cite {slug!r}: no paper in corpus — stub reference")
        return f"[missing paper {slug}] (cited slug not in corpus)"
    authors = _bibtex_authors(pref.authors).replace(" and ", "; ")
    # "Authors (year). Title." — robust plain-text assembly.
    head = " ".join(x for x in [authors, f"({pref.year})" if pref.year else ""] if x)
    line = (head + ". " if head else "") + (pref.title or slug) + "."
    try:
        alias = store.identifiers_for_refs([pref.id]).get(pref.id, {})
        if alias.get("doi"):
            line += f" doi:{alias['doi']}"
        elif alias.get("arxiv"):
            line += f" arXiv:{alias['arxiv']}"
    except Exception:
        pass
    return line


def _append_references(doc: Any, ctx: _Ctx) -> None:
    """A numbered **References** section: one entry per cited paper, in
    cite order, so entry ``n`` matches every in-text ``[n]`` mark. Each
    line is resolved through the SAME paper lookup as the ``.bib`` path
    (citation-integrity parity with the PDF). No-op when nothing was cited.

    A plain Word section (heading + numbered paragraphs), deliberately
    *not* native endnote fields: a paper cited many times must reuse one
    number, and a Word endnote can be referenced only once."""
    if not ctx.cited:
        return
    doc.add_heading("References", level=1)
    if ctx.endnote:
        # EndNote regenerates this list on "Update Citations and Bibliography"
        # from the records embedded in the in-text fields; the plain lines are
        # cached placeholder text shown until then.
        from precis.export.endnote import add_reflist_field

        cached = [
            f"[{i}] " + _format_reference(ctx.store, slug, ctx.warnings)
            for i, slug in enumerate(ctx.cited, start=1)
        ]
        add_reflist_field(doc, cached)
        return
    for i, slug in enumerate(ctx.cited, start=1):
        line = _format_reference(ctx.store, slug, ctx.warnings)
        p = doc.add_paragraph()
        p.add_run(f"[{i}] ").bold = True
        p.add_run(line)


def _append_acronyms(doc: Any, ctx: _Ctx) -> None:
    """An "Acronyms" list of every abbreviation actually used in the prose
    (auto-built, like the LaTeX glossaries acronym list) — SHORT → long."""
    used = sorted(s for s in ctx.used_acr if s in ctx.abbrevs)
    if not used:
        return
    doc.add_heading("Acronyms", level=1)
    for short in used:
        p = doc.add_paragraph()
        p.add_run(short).bold = True
        p.add_run(f" — {ctx.abbrevs[short]}")
